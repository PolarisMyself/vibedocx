"""Shared test fixtures."""
import os
import sys
import tempfile

import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))


@pytest.fixture
def tmp_docx():
    """Create a temporary .docx file with one paragraph."""
    tmp = os.path.join(tempfile.gettempdir(), 'test_tmp.docx')
    doc = Document()
    doc.add_paragraph('Test paragraph.')
    doc.save(tmp)
    yield tmp
    if os.path.exists(tmp):
        os.unlink(tmp)


@pytest.fixture
def blank_docx():
    """Create a blank temporary .docx file."""
    tmp = os.path.join(tempfile.gettempdir(), 'test_blank.docx')
    Document().save(tmp)
    yield tmp
    if os.path.exists(tmp):
        os.unlink(tmp)
