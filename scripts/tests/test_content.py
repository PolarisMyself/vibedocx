"""Tests for content.py operations."""
from docx import Document
from helper.content import append_content, page_break, replace_text, table_create


def test_append_heading(blank_docx):
    append_content(blank_docx, "heading", "Test Heading", level=1, output=blank_docx)
    doc = Document(blank_docx)
    texts = [p.text for p in doc.paragraphs]
    assert "Test Heading" in texts


def test_append_body(blank_docx):
    append_content(blank_docx, "body", "Sample body text.", output=blank_docx)
    doc = Document(blank_docx)
    assert any("Sample body text." in p.text for p in doc.paragraphs)


def test_replace_text(tmp_docx):
    replace_text(tmp_docx, {"Test": "Modified"}, output=tmp_docx)
    doc = Document(tmp_docx)
    assert "Modified paragraph" in doc.paragraphs[0].text


def test_replace_text_no_match(tmp_docx):
    count = replace_text(tmp_docx, {"NoSuchText": "X"}, output=tmp_docx)
    assert count == 0


def test_table_create(tmp_docx):
    table_create(tmp_docx, headers=["A", "B"], rows=[["1", "2"]],
                 caption="Table 1", output=tmp_docx)
    doc = Document(tmp_docx)
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "A"


def test_page_break(tmp_docx):
    count_before = len(list(Document(tmp_docx).paragraphs))
    page_break(tmp_docx, after_paragraph=0, output=tmp_docx)
    doc = Document(tmp_docx)
    # Page break attaches to existing paragraph, shouldn't increase count
    assert len(list(doc.paragraphs)) == count_before
