"""Tests for cite.py operations."""
from docx import Document
from helper.cite import ref_add, ref_cite, ref_list, ref_remove, ref_generate


def test_ref_add_and_cite_inline(blank_docx):
    # Add a paragraph for citing
    doc = Document(blank_docx)
    doc.add_paragraph('Citing test.')
    doc.save(blank_docx)

    ref_add(blank_docx, "test1", "Author. Title[J]. 2020.", output=blank_docx)
    ref_add(blank_docx, "test2", "Another. Book[M]. 2019.", output=blank_docx)
    ref_cite(blank_docx, "test1", 0, style='inline', output=blank_docx)

    info = ref_list(blank_docx)
    assert info['count'] == 2
    assert len(info['citations']) == 1


def test_ref_cite_footnote(blank_docx):
    doc = Document(blank_docx)
    doc.add_paragraph('Footnote test.')
    doc.save(blank_docx)

    ref_add(blank_docx, "fn1", "Author. Title[J]. 2020.", output=blank_docx)
    ref_cite(blank_docx, "fn1", 0, style='footnote', output=blank_docx)

    info = ref_list(blank_docx)
    assert len(info['citations']) == 1
    assert info['citations'][0]['style'] == 'footnote'


def test_ref_remove(blank_docx):
    doc = Document(blank_docx)
    doc.add_paragraph('Remove test.')
    doc.save(blank_docx)

    ref_add(blank_docx, "r1", "Author. Title[J]. 2020.", output=blank_docx)
    ref_cite(blank_docx, "r1", 0, style='inline', output=blank_docx)
    ref_remove(blank_docx, "r1", output=blank_docx)

    info = ref_list(blank_docx)
    assert info['count'] == 0


def test_ref_generate(blank_docx):
    ref_add(blank_docx, "g1", "Author. Title[J]. 2020.", output=blank_docx)
    ref_generate(blank_docx, heading="References", output=blank_docx)

    doc = Document(blank_docx)
    texts = [p.text for p in doc.paragraphs]
    assert any("Author. Title[J]" in t for t in texts)
