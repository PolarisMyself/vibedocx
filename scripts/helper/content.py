"""Content module — all content operations: replace, move, delete, insert.

Consolidates replace.py + move.py into a single content operations layer.
"""

import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from helper.filter import apply_filter, _find_heading_blocks
from helper.units import parse_size


def _make_t(text):
    """Create a w:t element with xml:space='preserve' — always use this."""
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    return t


# ──────────────────────────────────────────────
# Replace
# ──────────────────────────────────────────────

def replace_text(file_path, mapping, output=None):
    """Find and replace text in document (paragraphs, tables, headers/footers)."""
    doc = Document(file_path)
    count = 0
    for para in doc.paragraphs:
        count += _replace_in_paragraph(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _replace_in_paragraph(para, mapping)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    count += _replace_in_paragraph(para, mapping)
    doc.save(output)
    return count


def _replace_in_paragraph(para, mapping):
    """Replace text in a paragraph, handling cross-run cases."""
    count = 0
    for run in para.runs:
        for find_text, replace_with in mapping.items():
            if find_text in run.text:
                count += run.text.count(find_text)
                run.text = run.text.replace(find_text, replace_with)

    full_text = "".join(r.text for r in para.runs)
    needs_rebuild = any(ft in full_text for ft in mapping)
    if not needs_rebuild:
        return 0

    new_text = full_text
    for find_text, replace_with in mapping.items():
        if find_text in new_text:
            count += new_text.count(find_text)
            new_text = new_text.replace(find_text, replace_with)
    if new_text == full_text:
        return 0

    runs = para.runs
    total_len = len(new_text)
    original_total = sum(len(r.text) for r in runs)
    pos = 0
    for ri, run in enumerate(runs):
        if ri < len(runs) - 1:
            ratio = len(run.text) / original_total if original_total else 0
            run_len = round(len(run.text) + (total_len - original_total) * ratio)
            new_run_text = new_text[pos:pos + run_len]
        else:
            new_run_text = new_text[pos:]
        if new_run_text != run.text:
            run.text = new_run_text
        pos += len(new_run_text)
    return count


def replace_filtered(file_path, filter_spec, mapping, output=None):
    """Replace text only in elements matching a filter."""
    results = apply_filter(file_path, filter_spec)
    doc = Document(file_path)
    count = 0
    matched_runs = set()
    matched_paras = set()
    for r in results:
        if r.run is not None:
            matched_runs.add((r.paragraph_index, r.run_index))
        else:
            matched_paras.add(r.paragraph_index)
    for pi, para in enumerate(doc.paragraphs):
        if matched_paras and pi in matched_paras:
            for run in para.runs:
                for ft, rw in mapping.items():
                    if ft in run.text:
                        count += run.text.count(ft)
                        run.text = run.text.replace(ft, rw)
        if matched_runs:
            for ri, run in enumerate(para.runs):
                if (pi, ri) in matched_runs:
                    for ft, rw in mapping.items():
                        if ft in run.text:
                            count += run.text.count(ft)
                            run.text = run.text.replace(ft, rw)
    doc.save(output)
    return count


def replace_bookmark(file_path, bookmarks, output=None):
    """Replace content at bookmark locations."""
    doc = Document(file_path)
    count = 0
    body = doc.element.body
    bookmarks_to_process = []
    for bs in body.iter(qn("w:bookmarkStart")):
        name = bs.get(qn("w:name"))
        if name not in bookmarks:
            continue
        bm_id = bs.get(qn("w:id"))
        new_text = bookmarks[name]
        to_remove = []
        elem = bs
        while elem is not None:
            next_elem = elem.getnext()
            if elem.tag == qn("w:bookmarkEnd") and elem.get(qn("w:id")) == bm_id:
                break
            if elem is not bs:
                to_remove.append((elem.getparent(), elem))
            elem = next_elem
        for parent, child in to_remove:
            parent.remove(child)
        from docx.oxml import OxmlElement
        new_run = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_run.append(new_t)
        bs.addnext(new_run)
        count += 1
    doc.save(output)
    return count


# ──────────────────────────────────────────────
# Block operations (move / delete / swap)
# ──────────────────────────────────────────────

def _resolve_blocks(doc, filter_spec):
    """Resolve a filter spec to body child element blocks."""
    body = doc.element.body
    children = list(body)
    is_block = any(k in filter_spec for k in (
        "heading_block", "heading_block_regex", "heading_outline_level", "heading_level"))
    if is_block:
        results = _find_heading_blocks(doc, filter_spec)
        if not results:
            return []
        blocks = []
        for r in results:
            elems = [(children[ci], ci) for ci in range(r.block_start, r.block_end)]
            blocks.append(elems)
        return blocks
    if "paragraph_range" in filter_spec:
        start, end = filter_spec["paragraph_range"]
        para_idx = 0
        indices = []
        for bi, child in enumerate(children):
            tag = child.tag.split("}")[-1]
            if tag == "p":
                if start <= para_idx <= end:
                    indices.append(bi)
                para_idx += 1
        return [[(children[i], i) for i in indices]] if indices else []
    raise ValueError(f"Block filter required for move/delete/swap. Got: {filter_spec}")


def delete_block(file_path, filter_spec, output=None):
    """Delete a block from the document."""
    doc = Document(file_path)
    body = doc.element.body
    blocks = _resolve_blocks(doc, filter_spec)
    if not blocks:
        doc.save(output)
        return 0
    total = 0
    for block in reversed(blocks):
        for elem, _ in reversed(block):
            body.remove(elem)
            total += 1
    doc.save(output)
    return total


def move_block(file_path, filter_spec, target_filter, position="after", output=None):
    """Move a block to a new position."""
    doc = Document(file_path)
    body = doc.element.body
    src_blocks = _resolve_blocks(doc, filter_spec)
    if not src_blocks:
        doc.save(output)
        return 0
    tgt_blocks = _resolve_blocks(doc, target_filter)
    if not tgt_blocks:
        doc.save(output)
        return 0
    src_block = src_blocks[0]
    tgt_block = tgt_blocks[0]
    ref_elem = tgt_block[0][0] if position == "before" else tgt_block[-1][0]
    src_elems = [elem for elem, _ in src_block]
    for elem in src_elems:
        body.remove(elem)
    if position == "before":
        for elem in src_elems:
            ref_elem.addprevious(elem)
    else:
        for elem in reversed(src_elems):
            ref_elem.addnext(elem)
    doc.save(output)
    return len(src_elems)


def swap_blocks(file_path, filter1, filter2, output=None):
    """Swap two blocks atomically."""
    doc = Document(file_path)
    body = doc.element.body
    children = list(body)
    blocks1 = _resolve_blocks(doc, filter1)
    blocks2 = _resolve_blocks(doc, filter2)
    if not blocks1 or not blocks2:
        doc.save(output)
        return (0, 0)
    block_a = blocks1[0]
    block_b = blocks2[0]
    a_start = block_a[0][1]
    b_start = block_b[0][1]
    if a_start > b_start:
        block_a, block_b = block_b, block_a
        a_start, b_start = b_start, a_start
    elems_a = [elem for elem, _ in block_a]
    elems_b = [elem for elem, _ in block_b]
    for elem in reversed(elems_b):
        body.remove(elem)
    for elem in reversed(elems_a):
        body.remove(elem)
    current_children = list(body)
    if a_start < len(current_children):
        anchor = current_children[a_start]
        for elem in elems_b:
            anchor.addprevious(elem)
        last_b = elems_b[-1]
        for elem in reversed(elems_a):
            last_b.addnext(elem)
    else:
        for elem in elems_b:
            body.append(elem)
        for elem in elems_a:
            body.append(elem)
    doc.save(output)
    return (len(elems_a), len(elems_b))


# ──────────────────────────────────────────────
# Insert helpers
# ──────────────────────────────────────────────

def append_content(file_path, content_type, text, after_paragraph=None,
                   level=1, font_cn=None, font_en=None, size=None,
                   bold=None, italic=None, align=None, output=None):
    """Append a heading or body paragraph to the document.

    Args:
        file_path: Input .docx path.
        content_type: 'heading' or 'body'.
        text: Paragraph text content.
        after_paragraph: Insert after this paragraph index (None = append to end).
        level: Heading level (1-9), only for content_type='heading'.
        font_cn/font_en/size/bold/italic/align: Run/paragraph formatting.
        output: Output file path.
    """
    doc = Document(file_path)
    paras = list(doc.paragraphs)

    # Build the paragraph element
    p_elem = OxmlElement('w:p')

    if content_type == 'heading':
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), f'Heading{level}')
        pPr.append(pStyle)
        p_elem.append(pPr)

        r_elem = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = text
        r_elem.append(t_elem)
        _apply_run_fmt(r_elem, font_cn, font_en, size, bold, italic)
        p_elem.append(r_elem)

    else:  # body
        if align:
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), align)
            pPr.append(jc)
            p_elem.append(pPr)

        r_elem = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = text
        r_elem.append(t_elem)
        _apply_run_fmt(r_elem, font_cn or '宋体', font_en or 'Times New Roman', size or 12, bold, italic)
        p_elem.append(r_elem)

    # Insert at position
    if after_paragraph is not None:
        if after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range (0-{len(paras)-1})")
        ref_elem = paras[after_paragraph]._element
        ref_elem.addnext(p_elem)
    else:
        doc.element.body.append(p_elem)

    doc.save(output)
    return 1


def _apply_run_fmt(r_elem, font_cn=None, font_en=None, size=None, bold=None, italic=None):
    """Apply formatting properties to a w:r element."""
    if font_cn or font_en or size or bold is not None or italic is not None:
        rPr = OxmlElement('w:rPr')
        if font_cn or font_en:
            rFonts = OxmlElement('w:rFonts')
            if font_cn:
                rFonts.set(qn('w:eastAsia'), font_cn)
                rFonts.set(qn('w:ascii'), font_cn)
                rFonts.set(qn('w:hAnsi'), font_cn)
            if font_en:
                rFonts.set(qn('w:ascii'), font_en)
                rFonts.set(qn('w:hAnsi'), font_en)
            rPr.append(rFonts)
        if size is not None:
            pt_val = parse_size(size).pt
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(pt_val * 2)))
            rPr.append(sz)
        if bold is not None:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic is not None:
            i = OxmlElement('w:i')
            rPr.append(i)
        r_elem.insert(0, rPr)


def page_break(file_path, after_paragraph=None, output=None):
    """Insert a page break into the document.

    Attaches the break to an existing paragraph rather than creating a separate
    empty paragraph, avoiding blank-page artifacts.

    Args:
        file_path: Input .docx path.
        after_paragraph: Attach page break to this paragraph (None=last paragraph).
        output: Output file path.
    """
    doc = Document(file_path)
    paras = list(doc.paragraphs)

    if after_paragraph is not None:
        if after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range (0-{len(paras)-1})")
        target = paras[after_paragraph]
    elif paras:
        target = paras[-1]
    else:
        target = doc.add_paragraph()

    run = target.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    doc.save(output)
    return 1


# ──────────────────────────────────────────────
# Image insertion
# ──────────────────────────────────────────────

def insert_image(file_path, image_path, after_paragraph=None,
                 width=None, height=None, caption=None, ref_id=None, output=None):
    """Insert an image into the document.

    Args:
        file_path: Input .docx path.
        image_path: Path to image file (png, jpg, etc.).
        after_paragraph: Paragraph index to insert after (None = append at end).
        width: Width in cm (preserves aspect ratio if only width given).
        height: Height in cm.
        caption: Optional caption text (e.g. '系统架构图').
        ref_id: Optional reference ID for auto-numbering (e.g. 'fig:arch').
                When set, caption becomes '图X-Y caption' after numbering update.
        output: Output file path.
    """
    from docx.shared import Cm
    doc = Document(file_path)

    kwargs = {}
    if width:
        kwargs['width'] = Cm(width)
    if height:
        kwargs['height'] = Cm(height)
    if not kwargs:
        kwargs['width'] = Cm(12)

    # Always add at end first (python-docx needs properly attached paragraphs)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(image_path, **kwargs)

    # Add caption
    cap_para = None
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(4)
        cap_para.paragraph_format.space_after = Pt(8)
        cap_run = cap_para.add_run(caption)
        cap_run.font.size = Pt(10.5)

        # Store ref metadata and mark caption for auto-numbering
        if ref_id:
            from helper.numbering import mark_caption_with_ref, _load_figtab, _save_figtab
            mark_caption_with_ref(cap_run._element, ref_id, 'figure')
            figtab = _load_figtab(doc)
            figtab[ref_id] = {'type': 'figure', 'caption': caption}
            _save_figtab(doc, figtab)

    # Move to correct position if specified
    if after_paragraph is not None:
        paras = list(doc.paragraphs)
        if after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range")
        ref_elem = paras[after_paragraph]._element
        body = doc.element.body
        # Move image paragraph
        body.remove(img_para._element)
        ref_elem.addnext(img_para._element)
        # Move caption paragraph if exists
        if cap_para is not None:
            body.remove(cap_para._element)
            img_para._element.addnext(cap_para._element)

    doc.save(output)
    return 1


# ──────────────────────────────────────────────
# TOC (Table of Contents)
# ──────────────────────────────────────────────

def insert_toc(file_path, after_paragraph=None, levels=3, output=None):
    """Insert a Table of Contents field.

    Args:
        file_path: Input .docx path.
        after_paragraph: Insert after this paragraph index (None = append).
        levels: Number of heading levels to include (1-9).
        output: Output file path.
    """
    from docx.oxml import OxmlElement
    doc = Document(file_path)

    # Build TOC field
    # The TOC field code
    toc_instruction = f' TOC \\o "1-{levels}" \\h \\z \\u '

    # Create paragraph with TOC field
    if after_paragraph is not None:
        paras = list(doc.paragraphs)
        if after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range")
        ref_para = paras[after_paragraph]
        toc_p = OxmlElement('w:p')
        ref_para._element.addnext(toc_p)
    else:
        toc_p = doc.add_paragraph()._element

    # Add field begin
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    toc_p.append(r1)

    # Add field instruction
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = toc_instruction
    r2.append(instrText)
    toc_p.append(r2)

    # Add field separate
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    toc_p.append(r3)

    # Add placeholder text
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '（请在 Word 中右键更新目录）'
    r4.append(t)
    toc_p.append(r4)

    # Add field end
    r5 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r5.append(fldChar3)
    toc_p.append(r5)

    doc.save(output)
    return 1


# ──────────────────────────────────────────────
# Section management
# ──────────────────────────────────────────────

def insert_section(file_path, after_paragraph=None, break_type='nextPage',
                   header_text=None, footer_text=None, output=None):
    """Insert a section break with optional per-section header/footer.

    Args:
        file_path: Input .docx path.
        after_paragraph: Insert after this paragraph index.
        break_type: 'nextPage', 'continuous', 'evenPage', 'oddPage'.
        header_text: Header text for the new section.
        footer_text: Footer text for the new section.
        output: Output file path.
    """
    from docx.oxml import OxmlElement
    doc = Document(file_path)

    if after_paragraph is not None:
        paras = list(doc.paragraphs)
        if after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range")
        ref_para = paras[after_paragraph]

        # Insert section break paragraph after ref_para
        sect_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        sect_p.append(pPr)
        sectPr = OxmlElement('w:sectPr')

        # Section type
        type_elem = OxmlElement('w:type')
        type_elem.set(qn('w:val'), break_type)
        sectPr.append(type_elem)

        # Copy page setup from document
        doc_sect = doc.sections[0]
        pgSz = OxmlElement('w:pgSz')
        pgSz.set(qn('w:w'), str(int(doc_sect.page_width.emu)))
        pgSz.set(qn('w:h'), str(int(doc_sect.page_height.emu)))
        sectPr.append(pgSz)
        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'), str(int(doc_sect.top_margin.emu)))
        pgMar.set(qn('w:bottom'), str(int(doc_sect.bottom_margin.emu)))
        pgMar.set(qn('w:left'), str(int(doc_sect.left_margin.emu)))
        pgMar.set(qn('w:right'), str(int(doc_sect.right_margin.emu)))
        sectPr.append(pgMar)

        pPr.append(sectPr)
        ref_para._element.addnext(sect_p)

    doc.save(output)
    return 1


# ──────────────────────────────────────────────
# Table operations
# ──────────────────────────────────────────────

def table_add_row(file_path, table_index, data=None, output=None):
    """Add a row to a table."""
    doc = Document(file_path)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Table index {table_index} out of range (0-{len(tables)-1})")
    table = tables[table_index]
    row = table.add_row()
    if data:
        for i, text in enumerate(data):
            if i < len(row.cells):
                row.cells[i].text = str(text)
    doc.save(output)
    return 1


def table_add_column(file_path, table_index, header=None, width=None,
                     strategy="split", output=None):
    """Add a column to a table.

    Properly updates tblGrid (gridCol) and sets column width. For tables
    with merged cells, strategy controls behavior.

    Args:
        file_path: Input .docx path.
        table_index: 0-based table index.
        header: Optional header text for the new column.
        width: Column width in cm (float). Default None = auto-calculate
               from average of existing columns.
        strategy: For merged cells:
            "split" — Break gridSpan at the new column position (default).
            "expand" — Grow gridSpan of merged cells by 1 to include new column.
            "refuse" — Raise ValueError if any row has a merged cell at the
                       new column position.
    """
    from lxml import etree
    from docx.shared import Cm, Emu

    doc = Document(file_path)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Table index {table_index} out of range (0-{len(tables)-1})")
    table = tables[table_index]
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))
    gridCols = tbl.tblGrid.findall(qn('w:gridCol'))

    # Calculate default width from existing gridCols
    if width is not None:
        col_width = Cm(width)
    elif gridCols:
        total_emu = 0
        for gc in gridCols:
            w_val = gc.get(qn('w:w'))
            if w_val:
                total_emu += int(w_val)
        col_width = Emu(total_emu // max(len(gridCols), 1))
    else:
        col_width = Cm(2.5)

    # Add new gridCol (fixes the primary bug: gridCol was never added)
    new_gridCol = etree.SubElement(tbl.tblGrid, qn('w:gridCol'))
    new_gridCol.set(qn('w:w'), str(int(col_width.emu)))

    # Process each row
    new_col_pos = len(gridCols)  # new column is appended at the end

    for i, tr in enumerate(tr_list):
        tc_elements = tr.findall(qn('w:tc'))
        # Calculate grid offset for this row
        offset = 0
        for tc in tc_elements:
            tcPr = tc.find(qn('w:tcPr'))
            span = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    span = int(gs.get(qn('w:val')))
            tc_end = offset + span

            if offset <= new_col_pos < tc_end and i > 0:
                # New column intersects this merged cell
                if strategy == "refuse" and span > 1:
                    raise ValueError(
                        f"New column at position {new_col_pos} intersects a "
                        f"merged cell (gridSpan={span}) in row {i}. "
                        f"Use --strategy split or expand.")
                elif strategy == "split" and span > 1:
                    # Reduce gridSpan to make room for the new column
                    new_span = span - 1
                    if new_span <= 1:
                        if gs is not None:
                            tcPr.remove(gs)
                    else:
                        gs.set(qn('w:val'), str(new_span))
                elif strategy == "expand" and span > 1:
                    gs.set(qn('w:val'), str(span + 1))

            offset = tc_end

        # Append new tc to this row
        tc = etree.SubElement(tr, qn('w:tc'))
        tcPr_el = etree.SubElement(tc, qn('w:tcPr'))
        tcW = etree.SubElement(tcPr_el, qn('w:tcW'))
        tcW.set(qn('w:w'), str(int(col_width.emu)))
        tcW.set(qn('w:type'), 'dxa')
        p = etree.SubElement(tc, qn('w:p'))
        if header and i == 0:
            r = etree.SubElement(p, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = header

    doc.save(output)
    return 1


def table_merge_cells(file_path, table_index, row_start, col_start,
                      row_end, col_end, output=None):
    """Merge cells in a table."""
    doc = Document(file_path)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Table index {table_index} out of range")
    table = tables[table_index]
    cell_start = table.cell(row_start, col_start)
    cell_end = table.cell(row_end, col_end)
    cell_start.merge(cell_end)
    doc.save(output)
    return 1


def table_format_cell(file_path, table_index, row, col,
                      font_cn=None, font_en=None, size=None,
                      bold=None, italic=None,
                      align=None, vertical_align=None,
                      shading=None, width=None,
                      output=None):
    """Format a specific table cell.

    Args:
        file_path: Input .docx path.
        table_index: 0-based table index.
        row: 0-based row index.
        col: 0-based column index.
        font_cn: East Asian font name.
        font_en: Western font name.
        size: Font size (pt or Chinese name like "小四").
        bold: True/False/None (None = no change).
        italic: True/False/None.
        align: Horizontal alignment ("left", "center", "right", "justify").
        vertical_align: "top", "center", "bottom".
        shading: Hex fill color (e.g. "D9E2F3").
        width: Cell width in cm (float).
        output: Output file path.
    """
    from docx.shared import Cm
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from helper.config import ALIGN_MAP
    from helper.units import parse_size

    doc = Document(file_path)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Table index {table_index} out of range (0-{len(tables)-1})")
    table = tables[table_index]
    cell = table.cell(row, col)

    # Width
    if width is not None:
        cell.width = Cm(width)

    # Vertical alignment
    if vertical_align is not None:
        v_map = {"top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
                 "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                 "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM}
        va = v_map.get(vertical_align)
        if va is not None:
            cell.vertical_alignment = va

    # Shading
    if shading is not None:
        tcPr = cell._element.get_or_add_tcPr()
        existing = tcPr.find(qn('w:shd'))
        if existing is not None:
            tcPr.remove(existing)
        from lxml import etree
        shd = etree.SubElement(tcPr, qn('w:shd'))
        shd.set(qn('w:fill'), shading)
        shd.set(qn('w:val'), 'clear')

    # Font and alignment on cell paragraphs
    for para in cell.paragraphs:
        if align is not None:
            para.alignment = ALIGN_MAP.get(align)

        for run in para.runs:
            if font_cn or font_en:
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    from lxml import etree
                    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                if font_cn:
                    rFonts.set(qn('w:eastAsia'), font_cn)
                if font_en:
                    rFonts.set(qn('w:ascii'), font_en)
                    rFonts.set(qn('w:hAnsi'), font_en)
                if font_cn and not font_en:
                    rFonts.set(qn('w:ascii'), font_cn)
                    rFonts.set(qn('w:hAnsi'), font_cn)
            if size is not None:
                run.font.size = parse_size(size)
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

    doc.save(output)
    return 1


def table_delete_row(file_path, table_index, row_index, output=None):
    """Delete a row from a table. Content in merged cells cascades down.

    If the deleted row contains a vMerge="restart" cell, its content is
    moved to the next row's continue cell (which becomes the new restart).
    If the row contains only vMerge="continue" cells, the row is simply
    removed; the restart cell above keeps its content.

    Raises ValueError if the table has only one row.
    """
    from lxml import etree
    doc = Document(file_path)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Table index {table_index} out of range (0-{len(tables)-1})")
    table = tables[table_index]
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))
    if row_index < 0 or row_index >= len(tr_list):
        raise ValueError(f"Row index {row_index} out of range (0-{len(tr_list)-1})")
    if len(tr_list) <= 1:
        raise ValueError("Cannot delete the last row of a table; use table delete instead")

    tr = tr_list[row_index]
    tcs = tr.findall(qn('w:tc'))

    # Check for vMerge=restart cells and cascade content
    next_tr = tr_list[row_index + 1] if row_index + 1 < len(tr_list) else None
    for i, tc in enumerate(tcs):
        tcPr = tc.find(qn('w:tcPr'))
        vmerge_el = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
        vm_val = vmerge_el.get(qn('w:val')) if vmerge_el is not None else None

        if vm_val == 'restart' and next_tr is not None:
            # Find the corresponding continue cell in the next row
            next_tcs = next_tr.findall(qn('w:tc'))
            if i < len(next_tcs):
                next_tc = next_tcs[i]
                next_tcPr = next_tc.find(qn('w:tcPr'))
                next_vmerge = next_tcPr.find(qn('w:vMerge')) if next_tcPr is not None else None
                next_vm = next_vmerge.get(qn('w:val')) if next_vmerge is not None else None
                if next_vm == 'continue':
                    # Copy content from restart to continue cell
                    for child in list(next_tc):
                        if child.tag != qn('w:tcPr'):
                            next_tc.remove(child)
                    for child in list(tc):
                        if child.tag != qn('w:tcPr'):
                            tc.remove(child)
                            next_tc.append(child)
                    # Promote continue → restart
                    next_vmerge.set(qn('w:val'), 'restart')

    tbl.remove(tr)
    doc.save(output)
    return 1


def table_delete_column(file_path, table_index, col_index,
                        strategy="shrink", output=None):
    """Delete a column from a table.

    Args:
        file_path: Input .docx path.
        table_index: 0-based table index.
        col_index: 0-based grid column index to delete.
        strategy:
            "shrink" — cells spanning this column reduce gridSpan by 1 (default).
                       Content stays in the remaining (narrower) merged cell.
            "refuse" — raise ValueError if any cell intersecting this column
                       has gridSpan > 1 or vMerge set.

    Raises ValueError if the table has only one column with shrink strategy
    and no merged cells to absorb the deletion.
    """
    doc = Document(file_path)
    tables = doc.tables
    if table_index < 0 or table_index >= len(tables):
        raise ValueError(f"Table index {table_index} out of range (0-{len(tables)-1})")
    table = tables[table_index]
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))
    col_count = tbl.col_count

    if col_index < 0 or col_index >= col_count:
        raise ValueError(f"Column index {col_index} out of range (0-{col_count-1})")

    # Pre-scan for strategy validation
    for tr in tr_list:
        tc_elements = tr.findall(qn('w:tc'))
        offset = tr.grid_before
        for tc in tc_elements:
            tcPr = tc.find(qn('w:tcPr'))
            span = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    span = int(gs.get(qn('w:val')))
                vm = tcPr.find(qn('w:vMerge'))
                vm_val = vm.get(qn('w:val')) if vm is not None else None
            else:
                vm_val = None

            tc_end = offset + span
            if offset <= col_index < tc_end:
                if strategy == "refuse" and (span > 1 or vm_val is not None):
                    raise ValueError(
                        f"Column {col_index} intersects a merged cell "
                        f"(grid_span={span}, vMerge={vm_val}). Use --strategy shrink to shrink "
                        f"the merged region or handle manually.")
            offset = tc_end

    # Process each row
    for tr in tr_list:
        tc_elements = tr.findall(qn('w:tc'))
        offset = tr.grid_before
        for tc in list(tc_elements):
            tcPr = tc.find(qn('w:tcPr'))
            span = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    span = int(gs.get(qn('w:val')))
            tc_end = offset + span

            if offset <= col_index < tc_end:
                if strategy == "shrink":
                    if span > 1:
                        # Reduce gridSpan: n → n-1
                        new_span = span - 1
                        if new_span <= 1:
                            if tcPr is not None and gs is not None:
                                tcPr.remove(gs)
                        else:
                            if tcPr is None:
                                tcPr = tc.makeelement(qn('w:tcPr'), {})
                                tc.insert(0, tcPr)
                            if gs is None:
                                gs = tc.makeelement(qn('w:gridSpan'), {})
                                tcPr.append(gs)
                            gs.set(qn('w:val'), str(new_span))
                    else:
                        # span == 1: remove the cell entirely
                        tr.remove(tc)
                break  # only one cell per row spans this column
            offset = tc_end

    # Remove the gridCol
    gridCols = tbl.tblGrid.findall(qn('w:gridCol'))
    if col_index < len(gridCols):
        tbl.tblGrid.remove(gridCols[col_index])

    doc.save(output)
    return 1


def table_create(file_path, headers=None, rows=None, caption=None,
                 after_paragraph=None, ref_id=None, output=None):
    """Create a new table and insert into the document.

    Args:
        file_path: Input .docx path.
        headers: List of header cell texts.
        rows: List of lists, each inner list is a row of cell texts.
        caption: Optional caption text above the table.
        after_paragraph: Insert after this paragraph index (None = append).
        ref_id: Optional reference ID for auto-numbering (e.g. 'tab:compare').
        output: Output file path.
    """
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from lxml import etree

    doc = Document(file_path)
    headers = headers or []
    rows = rows or []

    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    num_rows = (1 if headers else 0) + len(rows)
    if num_cols == 0:
        doc.save(output)
        return 0

    # Collect elements to insert (caption + table)
    elements_to_insert = []

    # Caption paragraph
    if caption:
        cap_p = OxmlElement('w:p')
        cap_pPr = OxmlElement('w:pPr')
        cap_jc = OxmlElement('w:jc')
        cap_jc.set(qn('w:val'), 'center')
        cap_pPr.append(cap_jc)
        cap_spacing = OxmlElement('w:spacing')
        cap_spacing.set(qn('w:before'), '200')
        cap_spacing.set(qn('w:after'), '120')
        cap_pPr.append(cap_spacing)
        cap_p.append(cap_pPr)
        cap_r = OxmlElement('w:r')
        cap_rPr = OxmlElement('w:rPr')
        cap_b = OxmlElement('w:b')
        cap_rPr.append(cap_b)
        cap_rFonts = OxmlElement('w:rFonts')
        cap_rFonts.set(qn('w:eastAsia'), '黑体')
        cap_rFonts.set(qn('w:ascii'), 'Times New Roman')
        cap_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        cap_rPr.append(cap_rFonts)
        cap_sz = OxmlElement('w:sz')
        cap_sz.set(qn('w:val'), '21')
        cap_rPr.append(cap_sz)
        cap_r.append(cap_rPr)
        cap_t = OxmlElement('w:t')
        cap_t.text = caption
        cap_r.append(cap_t)
        cap_p.append(cap_r)
        elements_to_insert.append(cap_p)

        if ref_id:
            from helper.numbering import mark_caption_with_ref, _load_figtab, _save_figtab
            mark_caption_with_ref(cap_r, ref_id, 'table')
            figtab = _load_figtab(doc)
            figtab[ref_id] = {'type': 'table', 'caption': caption}
            _save_figtab(doc, figtab)

    # Build table via python-docx (for clean structure), then extract XML
    tmp_doc = Document()
    table = tmp_doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    borders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(borders, qn(f'w:{edge}'))
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')

    # Fill content
    row_idx = 0
    if headers:
        for j, text in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ''
            cp = cell.paragraphs[0]
            cp.alignment = 1  # CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(str(text))
            run.bold = True
            run.font.size = Pt(10.5)
        # Shade header row
        for cell in table.rows[0].cells:
            shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
            shading.set(qn('w:fill'), 'D9E2F3')
            shading.set(qn('w:val'), 'clear')
        row_idx = 1

    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            cell = table.rows[row_idx + i].cells[j]
            cell.text = ''
            cp = cell.paragraphs[0]
            cp.alignment = 1  # CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(str(text))
            run.font.size = Pt(10.5)

    elements_to_insert.append(table._tbl)

    # Insert at position
    if after_paragraph is not None:
        paras = list(doc.paragraphs)
        if after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range (0-{len(paras)-1})")
        ref_elem = paras[after_paragraph]._element
        for elem in reversed(elements_to_insert):
            ref_elem.addnext(elem)
    else:
        for elem in elements_to_insert:
            doc.element.body.append(elem)

    doc.save(output)
    return 1


# ──────────────────────────────────────────────
# Cross-references
# ──────────────────────────────────────────────

def xref_insert(file_path, ref_type, ref_id, text, paragraph_index,
                position='end', output=None):
    """Insert a cross-reference marker."""
    from lxml import etree
    doc = Document(file_path)
    paras = list(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= len(paras):
        raise ValueError(f"Paragraph index {paragraph_index} out of range")
    para = paras[paragraph_index]

    if position == 'end':
        run = para.add_run(text)
    else:
        from docx.oxml import OxmlElement
        run_elem = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        para._element.insert(0, run_elem)
        from docx.text.run import Run
        run = Run(run_elem, para)

    run.font.size = Pt(10.5)
    rPr = run._element.get_or_add_rPr()
    xref_mark = etree.SubElement(rPr, '{urn:opendocx:xref}xref')
    xref_mark.set('type', ref_type)
    xref_mark.set('id', ref_id)

    doc.save(output)
    return 1


def xref_update(file_path, labels, output=None):
    """Update cross-reference display text."""
    from lxml import etree
    doc = Document(file_path)
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                xref_mark = rPr.find('{urn:opendocx:xref}xref')
                if xref_mark is not None:
                    ref_id = xref_mark.get('id')
                    if ref_id in labels:
                        run.text = labels[ref_id]
                        count += 1
    doc.save(output)
    return count


# ──────────────────────────────────────────────
# Footnotes / Endnotes
# ──────────────────────────────────────────────

FOOTNOTES_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
)
FOOTNOTES_REL_TYPE = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
)
ENDNOTES_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml'
)
ENDNOTES_REL_TYPE = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
)

_NOTE_CONFIG = {
    'footnote': {
        'wrapper': 'footnotes', 'child': 'footnote',
        'ref_elem': 'footnoteReference', 'style': 'FootnoteReference',
        'content_type': FOOTNOTES_CONTENT_TYPE, 'rel_type': FOOTNOTES_REL_TYPE,
        'part_path': 'word/footnotes.xml', 'part_name': 'footnotes.xml',
    },
    'endnote': {
        'wrapper': 'endnotes', 'child': 'endnote',
        'ref_elem': 'endnoteReference', 'style': 'EndnoteReference',
        'content_type': ENDNOTES_CONTENT_TYPE, 'rel_type': ENDNOTES_REL_TYPE,
        'part_path': 'word/endnotes.xml', 'part_name': 'endnotes.xml',
    },
}


def _add_note_to_para(doc, para, text, note_type='footnote', after_text=None):
    """Add a footnote or endnote to a paragraph and return the note ID.

    Args:
        doc: python-docx Document (modified in place, NOT saved).
        para: Paragraph object to add the note reference to.
        text: Note content text.
        note_type: 'footnote' or 'endnote'.
        after_text: If set, insert after this text within the paragraph.

    Returns:
        The new note ID (int).
    """
    from lxml import etree
    cfg = _NOTE_CONFIG[note_type]

    # Find or create the notes part
    note_part = None
    for rel in doc.part.rels.values():
        if note_type in rel.reltype:
            note_part = rel.target_part
            break

    if note_part is not None:
        nx = etree.fromstring(note_part.blob)
        max_id = 0
        for child in nx.findall(qn(f'w:{cfg["child"]}')):
            fid = child.get(qn('w:id'))
            if fid and fid.lstrip('-').isdigit():
                max_id = max(max_id, int(fid))
        new_id = max(max_id + 1, 1)
    else:
        new_id = 1

    # Build the note reference run element
    ref_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), cfg['style'])
    rPr.append(rStyle)
    ref_run.append(rPr)
    note_ref = OxmlElement(f'w:{cfg["ref_elem"]}')
    note_ref.set(qn('w:id'), str(new_id))
    ref_run.append(note_ref)

    # Insert the reference run at the right position
    if after_text:
        _insert_element_after_text(para, after_text, ref_run)
    else:
        para._element.append(ref_run)

    # Add note content
    if note_part is not None:
        new_child = etree.SubElement(nx, qn(f'w:{cfg["child"]}'))
        new_child.set(qn('w:id'), str(new_id))
        n_para = etree.SubElement(new_child, qn('w:p'))
        n_pPr = etree.SubElement(n_para, qn('w:pPr'))
        n_spacing = etree.SubElement(n_pPr, qn('w:spacing'))
        n_spacing.set(qn('w:after'), '0')
        n_spacing.set(qn('w:line'), '240')
        n_spacing.set(qn('w:lineRule'), 'auto')
        n_ref_run = etree.SubElement(n_para, qn('w:r'))
        n_ref_rPr = etree.SubElement(n_ref_run, qn('w:rPr'))
        n_ref_style = etree.SubElement(n_ref_rPr, qn('w:rStyle'))
        n_ref_style.set(qn('w:val'), cfg['style'])
        n_text_run = etree.SubElement(n_para, qn('w:r'))
        n_t = etree.SubElement(n_text_run, qn('w:t'))
        n_t.set(qn('xml:space'), 'preserve')
        n_t.text = text
        note_part._blob = etree.tostring(nx, xml_declaration=True,
                                         encoding='UTF-8', standalone=True)
    else:
        attr = f'_vibedocx_pending_{note_type}s'
        setattr(doc, attr, {**getattr(doc, attr, {}), new_id: text})

    return new_id



def _insert_element_after_text(para, search_text, new_elem):
    """Insert an lxml element after a text pattern in a paragraph (handles cross-run)."""
    full = ''.join(r.text for r in para.runs)
    idx = full.find(search_text)
    if idx == -1:
        raise ValueError(f"Text '{search_text[:30]}...' not found in paragraph")

    end_pos = idx + len(search_text)
    pos = 0
    for ri, run in enumerate(para.runs):
        run_len = len(run.text)
        if pos + run_len >= end_pos:
            split_at = end_pos - pos
            if split_at < run_len:
                after_part = run.text[split_at:]
                run.text = run.text[:split_at]
                remainder = OxmlElement('w:r')
                orig_rPr = run._element.find(qn('w:rPr'))
                if orig_rPr is not None:
                    remainder.append(copy.deepcopy(orig_rPr))
                rt = OxmlElement('w:t')
                rt.set(qn('xml:space'), 'preserve')
                rt.text = after_part
                remainder.append(rt)
                run._element.addnext(remainder)
                target_elem = run._element
            else:
                target_elem = run._element
            target_elem.addnext(new_elem)
            return
        pos += run_len
    para._element.append(new_elem)


def insert_footnote(file_path, paragraph_index, text, output=None):
    """Insert a footnote at the end of a paragraph.

    Args:
        file_path: Input .docx path.
        paragraph_index: Paragraph to add footnote to.
        text: Footnote text.
        output: Output file path.
    """
    doc = Document(file_path)
    paras = list(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= len(paras):
        raise ValueError(f"Paragraph index {paragraph_index} out of range")
    para = paras[paragraph_index]

    new_id = _add_note_to_para(doc, para, text, 'footnote')
    doc.save(output)
    _inject_pending_notes(doc, output, 'footnote')
    return new_id


def _inject_pending_notes(doc, output_path, note_type):
    """If notes part was newly created, inject it into the docx zip."""
    pending = getattr(doc, f'_vibedocx_pending_{note_type}s', {})
    if not pending:
        return
    cfg = _NOTE_CONFIG[note_type]
    _inject_notes_part(output_path, pending, note_type, cfg)


def _inject_notes_part(output_path, note_texts, note_type, cfg):
    """Inject a notes part (footnotes or endnotes) into a saved docx zip."""
    import zipfile, os

    # Build minimal notes XML
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        f'<w:{cfg["wrapper"]} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        f'<w:{cfg["child"]} w:type="separator" w:id="-1">',
        '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f'<w:r><w:separator/></w:r></w:p>',
        f'</w:{cfg["child"]}>',
        f'<w:{cfg["child"]} w:type="continuationSeparator" w:id="0">',
        '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f'<w:r><w:continuationSeparator/></w:r></w:p>',
        f'</w:{cfg["child"]}>',
    ]
    for nid, ntext in note_texts.items():
        escaped = ntext.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        parts.append(
            f'<w:{cfg["child"]} w:id="{nid}">'
            '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            f'<w:r><w:rPr><w:rStyle w:val="{cfg["style"]}"/></w:rPr></w:r>'
            f'<w:r><w:t xml:space="preserve">{escaped}</w:t></w:r>'
            '</w:p>'
            f'</w:{cfg["child"]}>'
        )
    parts.append(f'</w:{cfg["wrapper"]}>')
    note_xml = ''.join(parts).encode('utf-8')

    tmp = output_path + '.tmp'
    with zipfile.ZipFile(output_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    ct = data.decode('utf-8')
                    if note_type not in ct:
                        ct = ct.replace(
                            '</Types>',
                            f'<Override PartName="/{cfg["part_path"]}" '
                            f'ContentType="{cfg["content_type"]}"/></Types>'
                        )
                    data = ct.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels = data.decode('utf-8')
                    if note_type not in rels:
                        rels = rels.replace(
                            '</Relationships>',
                            f'<Relationship Id="rId{note_type.title()}" '
                            f'Type="{cfg["rel_type"]}" '
                            f'Target="{cfg["part_name"]}"/></Relationships>'
                        )
                    data = rels.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr(cfg['part_path'], note_xml)
    os.replace(tmp, output_path)


# ─────── Note conversion ───────

def convert_notes(file_path, to_type, output=None):
    """Convert all footnotes to endnotes or vice versa.

    Args:
        file_path: Input .docx path.
        to_type: 'footnote' or 'endnote'.
        output: Output file path.
    """
    from lxml import etree

    from_type = 'endnote' if to_type == 'footnote' else 'footnote'
    src_cfg = _NOTE_CONFIG[from_type]
    dst_cfg = _NOTE_CONFIG[to_type]

    doc = Document(file_path)

    # ── 1. Convert references in body ──
    ref_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            r_elem = run._element

            # Replace reference element: footnoteReference ↔ endnoteReference
            old_ref = r_elem.find(qn(f'w:{src_cfg["ref_elem"]}'))
            if old_ref is None:
                continue
            note_id = old_ref.get(qn('w:id'))
            r_elem.remove(old_ref)
            new_ref = OxmlElement(f'w:{dst_cfg["ref_elem"]}')
            new_ref.set(qn('w:id'), note_id)
            r_elem.append(new_ref)

            # Replace run style
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is not None:
                rStyle = rPr.find(qn('w:rStyle'))
                if rStyle is not None and rStyle.get(qn('w:val')) == src_cfg['style']:
                    rStyle.set(qn('w:val'), dst_cfg['style'])

            ref_count += 1

    # ── 2. Move content between parts ──
    src_part = None
    for rel in doc.part.rels.values():
        if from_type in rel.reltype:
            src_part = rel.target_part
            break

    if src_part is not None:
        src_xml = etree.fromstring(src_part.blob)
        notes_data = []
        for child in src_xml.findall(qn(f'w:{src_cfg["child"]}')):
            child_id = child.get(qn('w:id'))
            if child_id and child_id.lstrip('-').isdigit() and int(child_id) > 0:
                # Extract note text
                note_text = ''
                for t_elem in child.iter(qn('w:t')):
                    note_text += (t_elem.text or '')
                notes_data.append((int(child_id), note_text))

        # Find or create destination part
        dst_part = None
        for rel in doc.part.rels.values():
            if to_type in rel.reltype:
                dst_part = rel.target_part
                break

        if dst_part is not None:
            dst_xml = etree.fromstring(dst_part.blob)
        else:
            empty_notes = (
                f'<w:{dst_cfg["wrapper"]} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                f'<w:{dst_cfg["child"]} w:type="separator" w:id="-1">'
                '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
                f'<w:r><w:separator/></w:r></w:p>'
                f'</w:{dst_cfg["child"]}>'
                f'<w:{dst_cfg["child"]} w:type="continuationSeparator" w:id="0">'
                '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
                f'<w:r><w:continuationSeparator/></w:r></w:p>'
                f'</w:{dst_cfg["child"]}>'
                f'</w:{dst_cfg["wrapper"]}>'
            )
            dst_xml = etree.fromstring(empty_notes)

        # Add notes to destination
        for nid, ntext in notes_data:
            new_child = etree.SubElement(dst_xml, qn(f'w:{dst_cfg["child"]}'))
            new_child.set(qn('w:id'), str(nid))
            n_para = etree.SubElement(new_child, qn('w:p'))
            n_pPr = etree.SubElement(n_para, qn('w:pPr'))
            n_spacing = etree.SubElement(n_pPr, qn('w:spacing'))
            n_spacing.set(qn('w:after'), '0')
            n_spacing.set(qn('w:line'), '240')
            n_spacing.set(qn('w:lineRule'), 'auto')
            n_ref_run = etree.SubElement(n_para, qn('w:r'))
            n_ref_rPr = etree.SubElement(n_ref_run, qn('w:rPr'))
            n_ref_style = etree.SubElement(n_ref_rPr, qn('w:rStyle'))
            n_ref_style.set(qn('w:val'), dst_cfg['style'])
            n_text_run = etree.SubElement(n_para, qn('w:r'))
            n_t = etree.SubElement(n_text_run, qn('w:t'))
            n_t.set(qn('xml:space'), 'preserve')
            n_t.text = ntext

        # Update or create destination part blob
        if dst_part is not None:
            dst_part._blob = etree.tostring(dst_xml, xml_declaration=True,
                                            encoding='UTF-8', standalone=True)
        else:
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI
            new_part = Part(
                PackURI(f'/{dst_cfg["part_path"]}'),
                dst_cfg['content_type'],
                etree.tostring(dst_xml, xml_declaration=True,
                              encoding='UTF-8', standalone=True),
                doc.part.package,
            )
            doc.part.relate_to(new_part, dst_cfg['rel_type'])

        # Clear source part (keep separator/continuation)
        for child in list(src_xml.findall(qn(f'w:{src_cfg["child"]}'))):
            child_id = child.get(qn('w:id'))
            if child_id and child_id.lstrip('-').isdigit() and int(child_id) > 0:
                src_xml.remove(child)
        src_part._blob = etree.tostring(src_xml, xml_declaration=True,
                                        encoding='UTF-8', standalone=True)

    doc.save(output)
    return ref_count


# ──────────────────────────────────────────────
# Equation (simple OMML)
# ──────────────────────────────────────────────

def insert_equation(file_path, paragraph_index, omml_xml, output=None):
    """Insert an OMML equation into a paragraph.

    Args:
        file_path: Input .docx path.
        paragraph_index: Target paragraph.
        omml_xml: OMML XML string for the equation.
        output: Output file path.
    """
    from lxml import etree
    doc = Document(file_path)
    paras = list(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= len(paras):
        raise ValueError(f"Paragraph index {paragraph_index} out of range")
    para = paras[paragraph_index]

    # Parse and insert the OMML element
    eq_elem = etree.fromstring(omml_xml)
    para._element.append(eq_elem)

    doc.save(output)
    return 1


# ──────────────────────────────────────────────
# Document merge
# ──────────────────────────────────────────────

def merge_documents(file_paths, output):
    """Merge multiple DOCX files into one.

    Copies body content, images, and other embedded resources from subsequent
    documents into the first.

    Args:
        file_paths: List of input .docx paths (in order).
        output: Output file path.
    """
    if not file_paths:
        raise ValueError("At least one input file required")

    from docx.oxml import OxmlElement
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    base_doc = Document(file_paths[0])

    # Track next rId number for new image relationships
    next_rid = _max_rid(base_doc) + 1
    next_image_idx = 1

    for fp in file_paths[1:]:
        other_doc = Document(fp)

        # Build rId mapping: old rId -> new rId for image references
        rid_map = {}
        for rId, rel in other_doc.part.rels.items():
            if 'image' not in rel.reltype:
                continue
            image_part = rel.target_part
            ext = (image_part.partname.ext or 'png').lstrip('.')
            new_name = f'image{next_image_idx}.{ext}'
            new_partname = f'/word/media/{new_name}'

            new_part = Part(
                PackURI(new_partname),
                image_part.content_type,
                image_part.blob,
                base_doc.part.package,
            )
            new_rId = base_doc.part.relate_to(
                new_part,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
            )
            rid_map[rId] = new_rId
            next_image_idx += 1

        # Add page break between documents
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        base_doc.element.body.append(p)

        # Copy body children, skipping sectPr, patching image references
        for child in list(other_doc.element.body):
            if child.tag == qn('w:sectPr'):
                continue
            child_copy = copy.deepcopy(child)
            _patch_image_refs(child_copy, rid_map)
            base_doc.element.body.append(child_copy)

    base_doc.save(output)
    return len(file_paths)


def _max_rid(doc):
    """Get the maximum numeric rId from document part relationships."""
    max_id = 0
    for rId in doc.part.rels:
        try:
            num = int(''.join(c for c in rId if c.isdigit()))
            if num > max_id:
                max_id = num
        except ValueError:
            pass
    return max_id


def _patch_image_refs(element, rid_map):
    """Replace r:embed attributes in a deep-copied XML element using a mapping."""
    if not rid_map:
        return
    r_embed = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    a_blip = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    for blip in element.iter(a_blip):
        old = blip.get(r_embed)
        if old in rid_map:
            blip.set(r_embed, rid_map[old])
