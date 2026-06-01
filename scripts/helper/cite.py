"""Reference management module - citations and bibliography generation.

References are stored as a custom element in document settings (word/settings.xml).
Two citation styles are supported:
  - inline: superscript [n] markers in body text (traditional)
  - footnote: Word native footnotes with auto-numbering (recommended)
"""

import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REFS_NS = 'urn:opendocx:refs'


# ── Storage helpers ──

def _load_refs(doc):
    """Load references from document settings."""
    settings = doc.settings.element
    refs_elem = settings.find(f'{{{REFS_NS}}}references')
    if refs_elem is not None and refs_elem.text:
        try:
            return json.loads(refs_elem.text)
        except json.JSONDecodeError:
            pass
    return {}


def _save_refs(doc, refs):
    """Save references to document settings."""
    from lxml import etree
    settings = doc.settings.element
    for old in settings.findall(f'{{{REFS_NS}}}references'):
        settings.remove(old)
    refs_elem = etree.SubElement(settings, f'{{{REFS_NS}}}references')
    refs_elem.text = json.dumps(refs, ensure_ascii=False)


def _load_fn_map(doc):
    """Load footnote-id → key mapping from settings."""
    settings = doc.settings.element
    elem = settings.find(f'{{{REFS_NS}}}fnmap')
    if elem is not None and elem.text:
        try:
            return json.loads(elem.text)
        except json.JSONDecodeError:
            pass
    return {}


def _save_fn_map(doc, fn_map):
    """Save footnote-id → key mapping to settings."""
    from lxml import etree
    settings = doc.settings.element
    for old in settings.findall(f'{{{REFS_NS}}}fnmap'):
        settings.remove(old)
    elem = etree.SubElement(settings, f'{{{REFS_NS}}}fnmap')
    elem.text = json.dumps(fn_map, ensure_ascii=False)


def _find_citation_runs(doc):
    """Find all citation runs in the document.

    Returns list of (paragraph, run, citation_number) tuples.
    """
    results = []
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                cite_mark = rPr.find(f'{{{REFS_NS}}}cite')
                if cite_mark is not None:
                    m = re.search(r'\[(\d+)\]', run.text or '')
                    if m:
                        results.append((para, run, int(m.group(1))))
    return results


# ── Public API ──

def ref_add(file_path, key, text, output=None):
    """Add a reference source to the document.

    key: unique identifier (e.g. 'heidegger1927')
    text: formatted reference text
    """
    doc = Document(file_path)
    refs = _load_refs(doc)
    refs[key] = text
    _save_refs(doc, refs)
    doc.save(output)
    return len(refs)


def ref_cite(file_path, key, paragraph_index, output=None, position='end',
             after_text=None, style='inline'):
    """Insert a citation in the specified paragraph.

    Two styles:
      inline   — superscript [n] marker with custom XML tag (traditional)
      footnote — Word native footnote, auto-numbered by Word

    key: reference key (must already be added via ref_add)
    paragraph_index: index of the paragraph to insert citation into
    position: 'end' or 'start' (only for style='inline' without after_text)
    after_text: insert citation right after this text within the paragraph
    style: 'inline' or 'footnote'
    """
    doc = Document(file_path)
    refs = _load_refs(doc)

    if key not in refs:
        raise ValueError(f"Reference '{key}' not found. Add it first with ref add.")

    ordered_keys = list(refs.keys())
    cite_num = ordered_keys.index(key) + 1

    paras = list(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= len(paras):
        raise ValueError(f"Paragraph index {paragraph_index} out of range (0-{len(paras)-1})")

    para = paras[paragraph_index]

    if style in ('footnote', 'endnote'):
        # Use Word native footnote/endnote — auto-numbered, proper superscript
        from helper.content import _add_note_to_para, _inject_pending_notes
        ref_text = refs[key]
        note_id = _add_note_to_para(doc, para, ref_text, style, after_text=after_text)
        # Track note_id → key for later management
        fn_map = _load_fn_map(doc)
        fn_map[str(note_id)] = key
        _save_fn_map(doc, fn_map)
        doc.save(output)
        _inject_pending_notes(doc, output, style)
        return cite_num

    # Inline style
    cite_text = f'[{cite_num}]'

    if after_text is not None:
        from helper.content import _insert_element_after_text
        # Build the citation run element first, then insert
        run_elem = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = cite_text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        _insert_element_after_text(para, after_text, run_elem)
        from docx.text.run import Run
        run = Run(run_elem, para)
    elif position == 'end':
        run = para.add_run(cite_text)
    else:
        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        run_elem.insert(0, rPr)
        t = OxmlElement('w:t')
        t.text = cite_text
        t.set(qn('xml:space'), 'preserve')
        run_elem.append(t)
        para._element.insert(0, run_elem)
        from docx.text.run import Run
        run = Run(run_elem, para)

    run.font.size = Pt(8)

    # Add superscript + citation mark
    rPr = run._element.get_or_add_rPr()
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    from lxml import etree
    cite_mark = etree.SubElement(rPr, f'{{{REFS_NS}}}cite')
    cite_mark.set('key', key)

    doc.save(output)
    return cite_num


def ref_list(file_path):
    """List all references and citations."""
    doc = Document(file_path)
    refs = _load_refs(doc)
    fn_map = _load_fn_map(doc)

    cite_info = []

    # Inline citations
    for para, run, num in _find_citation_runs(doc):
        rPr = run._element.find(qn('w:rPr'))
        key = None
        if rPr is not None:
            cite_mark = rPr.find(f'{{{REFS_NS}}}cite')
            if cite_mark is not None:
                key = cite_mark.get('key')
        cite_info.append({
            'number': num,
            'key': key,
            'style': 'inline',
            'text': run.text,
            'paragraph_text': para.text[:60],
        })

    # Footnote / Endnote citations
    for para in doc.paragraphs:
        for run in para.runs:
            for ref_tag in ('w:footnoteReference', 'w:endnoteReference'):
                fn_ref = run._element.find(qn(ref_tag))
                if fn_ref is not None:
                    fn_id = fn_ref.get(qn('w:id'))
                    if fn_id and fn_id in fn_map:
                        key = fn_map[fn_id]
                        style = 'footnote' if 'footnote' in ref_tag else 'endnote'
                        cite_info.append({
                            'number': int(fn_id),
                            'key': key,
                            'style': style,
                            'text': f'[{style} {fn_id}]',
                            'paragraph_text': para.text[:60],
                        })

    return {
        'references': refs,
        'citations': cite_info,
        'count': len(refs),
    }


def ref_generate(file_path, output=None, heading='参考文献',
                 style='numbered', after_heading=None):
    """Generate a bibliography section from stored references.

    If after_heading is specified, inserts bibliography after the first heading
    containing that text instead of appending to the end.
    """
    doc = Document(file_path)
    refs = _load_refs(doc)

    if not refs:
        raise ValueError("No references found. Add them first with ref add.")

    # Remove old bibliography
    _remove_old_bibliography(doc, heading)

    # Build bibliography paragraphs
    bib_paras = []

    # Heading
    from docx.oxml import OxmlElement
    h_para = OxmlElement('w:p')
    h_pPr = OxmlElement('w:pPr')
    h_pStyle = OxmlElement('w:pStyle')
    h_pStyle.set(qn('w:val'), 'Heading1')
    h_pPr.append(h_pStyle)
    h_para.append(h_pPr)
    h_run = OxmlElement('w:r')
    h_t = OxmlElement('w:t')
    h_t.text = heading
    h_run.append(h_t)
    h_para.append(h_run)
    bib_paras.append(h_para)

    # Reference entries
    for i, (key, text) in enumerate(refs.items(), 1):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        if style == 'numbered':
            t.text = f'[{i}] {text}'
        else:
            t.text = text
        r.append(t)
        p.append(r)
        bib_paras.append(p)

    # Insert bibliography
    if after_heading:
        inserted = False
        body = doc.element.body
        children = list(body)
        for ci, child in enumerate(children):
            if child.tag != qn('w:p'):
                continue
            text = ''
            for t_elem in child.iter(qn('w:t')):
                text += (t_elem.text or '')
            if after_heading in text:
                # Check if this is a heading
                pPr = child.find(qn('w:pPr'))
                is_heading = False
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        val = pStyle.get(qn('w:val'), '')
                        if val.lower().startswith('heading'):
                            is_heading = True
                if is_heading:
                    for bib_p in reversed(bib_paras):
                        child.addnext(bib_p)
                    inserted = True
                    break
        if not inserted:
            for bib_p in bib_paras:
                body.append(bib_p)
    else:
        for bib_p in bib_paras:
            doc.element.body.append(bib_p)

    doc.save(output)
    return len(refs)


def _remove_old_bibliography(doc, heading):
    """Remove existing bibliography section if present."""
    body = doc.element.body
    children = list(body)
    in_bib = False
    to_remove = []

    for child in children:
        if child.tag == qn('w:p'):
            text = ''
            for t_elem in child.iter(qn('w:t')):
                text += (t_elem.text or '')

            pPr = child.find(qn('w:pPr'))
            is_h1 = False
            if pPr is not None:
                style_el = pPr.find(qn('w:pStyle'))
                if style_el is not None:
                    val = style_el.get(qn('w:val'), '')
                    if val.lower() in ('heading1', 'heading 1'):
                        is_h1 = True
                outline = pPr.find(qn('w:outlineLvl'))
                if outline is not None and outline.get(qn('w:val')) == '0':
                    is_h1 = True

            if is_h1 and heading in text:
                in_bib = True
                to_remove.append(child)
                continue

            if in_bib and is_h1 and heading not in text:
                break

        if in_bib:
            to_remove.append(child)

    for elem in to_remove:
        body.remove(elem)


def ref_renumber(file_path, output=None):
    """Renumber all citations based on current reference order."""
    doc = Document(file_path)
    refs = _load_refs(doc)
    ordered_keys = list(refs.keys())
    key_to_num = {key: i + 1 for i, key in enumerate(ordered_keys)}

    count = 0
    for para, run, old_num in _find_citation_runs(doc):
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            cite_mark = rPr.find(f'{{{REFS_NS}}}cite')
            if cite_mark is not None:
                key = cite_mark.get('key')
                if key in key_to_num:
                    new_num = key_to_num[key]
                    if new_num != old_num:
                        run.text = f'[{new_num}]'
                        count += 1

    doc.save(output)
    return count


def ref_remove(file_path, key, output=None):
    """Remove a reference and all its citations (inline, footnote, and endnote)."""
    from lxml import etree

    doc = Document(file_path)
    refs = _load_refs(doc)

    if key not in refs:
        raise ValueError(f"Reference '{key}' not found.")

    del refs[key]
    _save_refs(doc, refs)

    fn_map = _load_fn_map(doc)

    # Remove inline citation runs for this key
    for para in doc.paragraphs:
        for run in list(para.runs):
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                cite_mark = rPr.find(f'{{{REFS_NS}}}cite')
                if cite_mark is not None and cite_mark.get('key') == key:
                    para._element.remove(run._element)

    # Remove footnote/endnote references and content
    fn_ids_to_remove = [fid for fid, k in fn_map.items() if k == key]
    for fn_id in fn_ids_to_remove:
        # Remove body reference runs
        for para in doc.paragraphs:
            for run in list(para.runs):
                for ref_tag in ('w:footnoteReference', 'w:endnoteReference'):
                    fn_ref = run._element.find(qn(ref_tag))
                    if fn_ref is not None and fn_ref.get(qn('w:id')) == fn_id:
                        para._element.remove(run._element)

        # Remove note content from footnotes/endnotes parts
        for note_type in ('footnote', 'endnote'):
            for rel in doc.part.rels.values():
                if note_type in rel.reltype:
                    note_part = rel.target_part
                    nx = etree.fromstring(note_part.blob)
                    for child in nx.findall(qn(f'w:{note_type}')):
                        if child.get(qn('w:id')) == fn_id:
                            nx.remove(child)
                            break
                    note_part._blob = etree.tostring(nx, xml_declaration=True,
                                                     encoding='UTF-8', standalone=True)

        # Remove from fn_map
        del fn_map[fn_id]

    _save_fn_map(doc, fn_map)
    doc.save(output)

    # Renumber inline citations in the saved file
    if output:
        ref_renumber(output, output=output)

    return len(refs)
