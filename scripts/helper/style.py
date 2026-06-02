"""Format module - modify styles and clear direct formatting."""

import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from helper.config import ALIGN_MAP, ALIGN_REVERSE, PAPER_SIZES
from helper.units import parse_size, parse_indent


def format_style(file_path, style_name, font=None, font_west=None, size=None,
                 bold=False, no_bold=False, italic=False, no_italic=False,
                 color=None, no_color=False,
                 align=None, line_spacing=None, space_before=None, space_after=None,
                 indent_first=None, output=None):
    """Modify a named style's properties. Saves to output path."""
    doc = Document(file_path)

    try:
        style = doc.styles[style_name]
    except KeyError:
        raise ValueError(f"Style '{style_name}' not found")

    # Font properties
    if font is not None:
        # Set East Asian font via XML (python-docx doesn't have direct API)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), font)
        # Also set the general font name
        style.font.name = font

    if font_west is not None:
        style.font.name = font_west

    if size is not None:
        style.font.size = parse_size(size)

    if bold:
        style.font.bold = True
    elif no_bold:
        style.font.bold = False

    if italic:
        style.font.italic = True
    elif no_italic:
        style.font.italic = False

    # Color (w:color on rPr -- not exposed by python-docx style.font)
    if color is not None or no_color:
        rpr_c = style.element.get_or_add_rPr()
        existing_colors = rpr_c.findall(qn("w:color"))
        for ec in existing_colors:
            rpr_c.remove(ec)
        if color is not None and not no_color:
            color_el = rpr_c.makeelement(qn("w:color"), {})
            color_el.set(qn("w:val"), color)
            rpr_c.append(color_el)

    # Paragraph properties (only for paragraph styles)
    if style.type == 1:  # paragraph style
        pf = style.paragraph_format

        if align is not None:
            pf.alignment = ALIGN_MAP.get(align)
        if line_spacing is not None:
            pf.line_spacing = line_spacing
        if space_before is not None:
            pf.space_before = Pt(space_before)
        if space_after is not None:
            pf.space_after = Pt(space_after)
        if indent_first is not None:
            fs = size if size is not None else (style.font.size.pt if style.font.size else 12)
            pf.first_line_indent = parse_indent(indent_first, fs)

    doc.save(output)


def format_clear_direct(file_path, range_start=None, range_end=None,
                        style_filter=None, output=None):
    """Clear direct formatting from paragraphs, reverting to style defaults."""
    doc = Document(file_path)
    all_paras = list(doc.paragraphs)

    # Determine range
    start = range_start or 0
    end = range_end if range_end is not None else len(all_paras) - 1
    start = max(0, start)
    end = min(len(all_paras) - 1, end)

    for i in range(start, end + 1):
        para = all_paras[i]

        # Style filter
        if style_filter and para.style.name != style_filter:
            continue

        # Clear direct paragraph formatting (pPr)
        ppr = para._element.find(qn("w:pPr"))
        if ppr is not None:
            for tag in ["w:jc", "w:spacing", "w:ind", "w:pBdr", "w:shd",
                        "w:numPr", "w:tabs", "w:keepNext", "w:keepLines",
                        "w:pageBreakBefore", "w:widowControl", "w:autoSpaceDE",
                        "w:autoSpaceDN", "w:textAlignment", "w:textDirection"]:
                for elem in ppr.findall(qn(tag)):
                    ppr.remove(elem)

        # Clear direct run formatting (rPr) — preserve custom XML markers
        for run in para.runs:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                for child in list(rpr):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ('b', 'bCs', 'i', 'iCs', 'u', 'sz', 'szCs',
                               'rFonts', 'color', 'highlight', 'vertAlign',
                               'spacing', 'position', 'strike', 'dstrike',
                               'outline', 'shadow', 'emboss', 'imprint',
                               'smallCaps', 'caps', 'vanish', 'specVanish',
                               'webHidden', 'lang', 'effect'):
                        rpr.remove(child)

    doc.save(output)


def format_paragraph(file_path, paragraph_index, font=None, font_west=None,
                     size=None, bold=None, italic=None, align=None,
                     line_spacing=None, output=None):
    """Apply direct formatting to a specific paragraph (does not modify style).

    Args:
        file_path: Input .docx path.
        paragraph_index: Index of paragraph to format.
        font/font_west/size/bold/italic: Run-level formatting.
        align/line_spacing: Paragraph-level formatting.
        output: Output file path.
    """
    from docx.oxml import OxmlElement
    doc = Document(file_path)
    paras = list(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= len(paras):
        raise ValueError(f"Paragraph index {paragraph_index} out of range (0-{len(paras)-1})")
    para = paras[paragraph_index]

    # Paragraph-level formatting
    pPr = para._element.get_or_add_pPr()
    if align is not None:
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), align)
    if line_spacing is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
        spacing.set(qn('w:lineRule'), 'auto')

    # Run-level formatting
    for run in para.runs:
        rPr = run._element.get_or_add_rPr()
        if font:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), font)
            rFonts.set(qn('w:ascii'), font)
            rFonts.set(qn('w:hAnsi'), font)
        if font_west:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), font_west)
            rFonts.set(qn('w:hAnsi'), font_west)
        if size is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            pt_val = parse_size(size).pt
            sz.set(qn('w:val'), str(int(pt_val * 2)))
        if bold is not None:
            b = rPr.find(qn('w:b'))
            if b is not None:
                rPr.remove(b)
            if bold:
                rPr.append(OxmlElement('w:b'))
        if italic is not None:
            i = rPr.find(qn('w:i'))
            if i is not None:
                rPr.remove(i)
            if italic:
                rPr.append(OxmlElement('w:i'))

    doc.save(output)


def format_page(file_path, margin_top=None, margin_bottom=None,
                margin_left=None, margin_right=None, paper=None,
                orientation=None, header_text=None, footer_text=None,
                output=None):
    """Set page properties (margins, paper size, orientation, header/footer).

    Margins in cm. paper: "A4", "A3", "Letter". orientation: "portrait" or "landscape".
    """
    doc = Document(file_path)
    section = doc.sections[0]

    # Margins (cm)
    if margin_top is not None:
        section.top_margin = Cm(margin_top)
    if margin_bottom is not None:
        section.bottom_margin = Cm(margin_bottom)
    if margin_left is not None:
        section.left_margin = Cm(margin_left)
    if margin_right is not None:
        section.right_margin = Cm(margin_right)

    # Paper size
    PAPER_SIZES = {
        "A4": (21.0, 29.7),
        "A3": (29.7, 42.0),
        "Letter": (21.59, 27.94),
    }
    if paper is not None:
        w, h = PAPER_SIZES.get(paper.upper(), (21.0, 29.7))
        section.page_width = Cm(w)
        section.page_height = Cm(h)

    # Orientation
    if orientation is not None:
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            if section.page_width < section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
        elif orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width

    # Header
    if header_text is not None:
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            header.paragraphs[0].text = header_text
        else:
            header.add_paragraph(header_text)

    # Footer
    if footer_text is not None:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            footer.paragraphs[0].text = footer_text
        else:
            footer.add_paragraph(footer_text)

    doc.save(output)


def style_export(file_path, output_json):
    """Export style definitions to a JSON file.

    Args:
        file_path: Input .docx path.
        output_json: Output JSON file path.

    Returns:
        Number of styles exported.
    """
    import json
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    styles = []
    for style in doc.styles:
        if style.type not in (1, 2):
            continue
        info = {
            "name": style.name,
            "type": "paragraph" if style.type == 1 else "character",
        }
        font = style.font
        if font.name: info["font"] = font.name
        if font.size: info["size"] = font.size.pt
        if font.bold is not None: info["bold"] = font.bold
        if font.italic is not None: info["italic"] = font.italic

        rpr = style.element.find(qn("w:rPr"))
        if rpr is not None:
            # Color
            color_el = rpr.find(qn("w:color"))
            if color_el is not None:
                info["color_rgb"] = color_el.get(qn("w:val"))
                theme = color_el.get(qn("w:themeColor"))
                if theme:
                    info["color_theme"] = theme
            # Fonts
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                ea = rfonts.get(qn("w:eastAsia"))
                if ea: info["font_cn"] = ea
                ascii_font = rfonts.get(qn("w:ascii"))
                if ascii_font: info["font_ascii"] = ascii_font

        if style.type == 1:
            pf = style.paragraph_format
            if pf.alignment: info["align"] = ALIGN_REVERSE.get(pf.alignment)
            if pf.line_spacing: info["line_spacing"] = pf.line_spacing
            if pf.space_before: info["space_before"] = pf.space_before.pt
            if pf.space_after: info["space_after"] = pf.space_after.pt
            if pf.first_line_indent: info["indent_first_pt"] = pf.first_line_indent.pt

        styles.append(info)

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(styles, f, ensure_ascii=False, indent=2)
    return len(styles)


def style_import(file_path, json_path, output):
    """Import style definitions from a JSON file.

    Args:
        file_path: Input .docx path.
        json_path: JSON file with style definitions.
        output: Output file path.
    """
    import json
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    with open(json_path, 'r', encoding='utf-8') as f:
        styles = json.load(f)

    doc = Document(file_path)
    count = 0
    for sdef in styles:
        name = sdef.get("name")
        if not name:
            continue
        try:
            style = doc.styles[name]
        except KeyError:
            continue

        if "font" in sdef:
            style.font.name = sdef["font"]
        if "font_cn" in sdef:
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                from lxml import etree
                rfonts = etree.SubElement(rpr, qn("w:rFonts"))
            rfonts.set(qn("w:eastAsia"), sdef["font_cn"])
        if "size" in sdef:
            style.font.size = parse_size(sdef["size"])
        if "bold" in sdef:
            style.font.bold = sdef["bold"]
        if "italic" in sdef:
            style.font.italic = sdef["italic"]

        # Color import
        if "color_rgb" in sdef:
            rpr = style.element.get_or_add_rPr()
            for existing in rpr.findall(qn("w:color")):
                rpr.remove(existing)
            color_el = rpr.makeelement(qn("w:color"), {})
            color_el.set(qn("w:val"), sdef["color_rgb"])
            if "color_theme" in sdef:
                color_el.set(qn("w:themeColor"), sdef["color_theme"])
            rpr.append(color_el)

        if style.type == 1:
            pf = style.paragraph_format
            if "align" in sdef:
                pf.alignment = ALIGN_MAP.get(sdef["align"])
            if "line_spacing" in sdef:
                pf.line_spacing = sdef["line_spacing"]
            if "space_before" in sdef:
                pf.space_before = Pt(sdef["space_before"])
            if "space_after" in sdef:
                pf.space_after = Pt(sdef["space_after"])
            if "indent_first_pt" in sdef:
                pf.first_line_indent = Pt(sdef["indent_first_pt"])

        count += 1

    doc.save(output)
    return count
