"""Query module — read-only document inspection and element selection.

Merges inspect.py + select.py into a single query layer.
"""

import os
from docx import Document
from docx.oxml.ns import qn
from helper.filter import apply_filter, get_outline_level


# ──────────────────────────────────────────────
# Inspect
# ──────────────────────────────────────────────



def inspect_structure(file_path):
    """Return document structure overview."""
    doc = Document(file_path)
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name or ""
        outline_lvl = get_outline_level(para)
        is_heading = style_name.startswith("Heading") or (outline_lvl is not None)
        if not is_heading:
            continue
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 0
        else:
            level = outline_lvl + 1 if outline_lvl is not None else 0
        info = {"index": i, "level": level, "text": para.text, "style": style_name}
        if outline_lvl is not None:
            info["outline_level"] = outline_lvl
        headings.append(info)

    image_count = sum(1 for _ in doc.element.body.iter(qn("a:blip")))
    return {
        "filename": os.path.basename(file_path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": image_count,
        "sections": len(doc.sections),
        "headings": headings,
    }


def inspect_styles(file_path, style_name=None):
    """Return style definitions."""
    from helper.config import ALIGN_REVERSE as _align_map
    doc = Document(file_path)
    styles = []
    for style in doc.styles:
        if style.type not in (1, 2):
            continue
        if style_name and style.name != style_name:
            continue
        info = {"name": style.name, "type": "paragraph" if style.type == 1 else "character"}
        font = style.font
        info["font"] = font.name
        info["font_size_pt"] = font.size.pt if font.size else None
        info["bold"] = font.bold
        info["italic"] = font.italic
        rpr = style.element.find(qn("w:rPr"))
        if rpr is not None:
            color_el = rpr.find(qn("w:color"))
            if color_el is not None:
                info["color_rgb"] = color_el.get(qn("w:val"))
                info["color_theme"] = color_el.get(qn("w:themeColor"))
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                info["font_ascii"] = rfonts.get(qn("w:ascii"))
                info["font_hAnsi"] = rfonts.get(qn("w:hAnsi"))
                info["font_cn"] = rfonts.get(qn("w:eastAsia"))
                ea = rfonts.get(qn("w:eastAsia"))
                if ea:
                    info["font"] = ea
        if style.type == 1:
            pf = style.paragraph_format
            info["alignment"] = _align_map.get(pf.alignment, str(pf.alignment) if pf.alignment else None)
            info["line_spacing"] = pf.line_spacing
            info["space_before_pt"] = pf.space_before.pt if pf.space_before else None
            info["space_after_pt"] = pf.space_after.pt if pf.space_after else None
            fli = pf.first_line_indent
            if fli is not None:
                fs = font.size.pt if font.size else 12
                info["first_line_indent_chars"] = round(fli.pt / (fs * 0.35), 1) if fs else None
            else:
                info["first_line_indent_chars"] = None
        styles.append(info)
    return {"styles": styles}


def inspect_formatting(file_path, range_start=None, range_end=None, paragraph_index=None):
    """Return paragraph formatting details."""
    doc = Document(file_path)
    all_paras = list(doc.paragraphs)
    if paragraph_index is not None:
        if paragraph_index < 0 or paragraph_index >= len(all_paras):
            return {"paragraphs": []}
        paras = [(paragraph_index, all_paras[paragraph_index])]
    else:
        start = max(0, range_start or 0)
        end = min(len(all_paras) - 1, range_end if range_end is not None else len(all_paras) - 1)
        paras = [(i, all_paras[i]) for i in range(start, end + 1)] if start <= end else []
    result = []
    for idx, para in paras:
        pf = para.paragraph_format
        info = {
            "index": idx, "text": para.text,
            "style_name": para.style.name if para.style else None,
            "paragraph_format": {
                "alignment": str(pf.alignment) if pf.alignment else None,
                "line_spacing": pf.line_spacing,
                "first_line_indent_chars": (
                    round(pf.first_line_indent.pt / 12, 1)
                    if pf.first_line_indent and pf.first_line_indent.pt else None
                ),
            },
        }
        ppr = para._element.find(qn("w:pPr"))
        direct_fmt = {}
        if ppr is not None:
            jc = ppr.find(qn("w:jc"))
            if jc is not None:
                direct_fmt["alignment"] = jc.get(qn("w:val"))
            spacing = ppr.find(qn("w:spacing"))
            if spacing is not None:
                line_val = spacing.get(qn("w:line"))
                if line_val:
                    direct_fmt["line_spacing"] = int(line_val) / 240
        info["direct_formatting"] = direct_fmt
        runs = []
        for run in para.runs:
            ri = {"text": run.text}
            if run.bold: ri["bold"] = True
            if run.italic: ri["italic"] = True
            if run.font.name: ri["font"] = run.font.name
            if run.font.size: ri["size_pt"] = run.font.size.pt
            runs.append(ri)
        info["runs"] = runs
        result.append(info)
    return {"paragraphs": result}


def inspect_content(file_path, range_start=None, range_end=None, formatted=False):
    """Return document content as text with position info."""
    doc = Document(file_path)
    all_paras = list(doc.paragraphs)
    total = len(all_paras)
    start = max(0, range_start or 0)
    end = min(total - 1, range_end if range_end is not None else total - 1)
    if start > end:
        return {"total": total, "paragraphs": []}
    result = []
    for i in range(start, end + 1):
        para = all_paras[i]
        if not formatted:
            result.append({"index": i, "text": para.text})
        else:
            style_name = para.style.name if para.style else ""
            outline_lvl = get_outline_level(para)
            entry = {"index": i, "text": para.text, "style": style_name}
            if outline_lvl is not None:
                entry["outline_level"] = outline_lvl
            if any(r.bold or r.italic for r in para.runs):
                entry["runs"] = [
                    {"text": r.text, **({"bold": True} if r.bold else {}),
                     **({"italic": True} if r.italic else {}),
                     **({"font": r.font.name} if r.font.name else {}),
                     **({"size": r.font.size.pt} if r.font.size else {})}
                    for r in para.runs
                ]
            result.append(entry)
    return {"total": total, "range": [start, end], "paragraphs": result}


# ──────────────────────────────────────────────
# Select
# ──────────────────────────────────────────────

def select(file_path, filter_spec, context_lines=0):
    """Select elements matching a filter."""
    results = apply_filter(file_path, filter_spec)
    doc = Document(file_path)
    all_paras = list(doc.paragraphs)
    matches = []
    for r in results:
        m = r.to_dict()
        if r.element_type == "run" and r.run:
            fmt = {}
            if r.run.bold: fmt["bold"] = True
            if r.run.italic: fmt["italic"] = True
            if r.run.font.name: fmt["font"] = r.run.font.name
            if r.run.font.size: fmt["size_pt"] = r.run.font.size.pt
            m["run_format"] = fmt
        if r.paragraph:
            pfmt = {}
            if r.paragraph.style: pfmt["style"] = r.paragraph.style.name
            pf = r.paragraph.paragraph_format
            if pf.alignment: pfmt["alignment"] = str(pf.alignment)
            if pf.line_spacing: pfmt["line_spacing"] = pf.line_spacing
            m["para_format"] = pfmt
            outline_lvl = get_outline_level(r.paragraph)
            if outline_lvl is not None:
                m["outline_level"] = outline_lvl
        if context_lines > 0 and r.paragraph is not None:
            start = max(0, r.paragraph_index - context_lines)
            end = min(len(all_paras) - 1, r.paragraph_index + context_lines)
            m["context"] = [
                {"index": ci, "text": all_paras[ci].text, "is_match": ci == r.paragraph_index}
                for ci in range(start, end + 1)
            ]
        matches.append(m)
    return {"total_matches": len(matches), "matches": matches}
