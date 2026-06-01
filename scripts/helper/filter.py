"""Filter engine - select document elements by composable criteria.

Filter spec format:
  - dict/object  → AND (all fields must match)
  - list/array   → OR (any element must match)
  - {"and": [...]} → explicit AND
  - {"or":  [...]} → explicit OR

Leaf filter fields (on paragraphs):
  - style: str              - style name (exact match)
  - outline_level: int      - outline level (0=chapter, 1=section, 2=subsection, ...)
  - text_contains: str      - paragraph text contains substring
  - text_regex: str         - paragraph text matches regex
  - bold: bool              - paragraph has any bold run
  - italic: bool            - paragraph has any italic run
  - font: str               - any run uses this font
  - paragraph_range: [int, int] - index in [start, end] inclusive
  - in_table: bool          - paragraph is inside a table cell
  - under_heading: str      - under nearest heading with this text (contains match)

Leaf filter fields (on runs):
  - run_text_contains: str  - run text contains substring
  - run_text_regex: str     - run text matches regex
  - run_bold: bool          - run is bold
  - run_italic: bool        - run is italic
  - run_font: str           - run font name
  - run_size: float         - run font size in pt

Block filter fields (returns body child ranges, not paragraphs):
  - heading_block: str      - select entire block under this heading (contains match)
  - heading_block_regex: str - regex version of heading_block
  - heading_level: int      - heading level for heading_block (default: auto-detect)
  - heading_outline_level: int - outline level for heading_block (default: auto-detect)
"""

import re
from dataclasses import dataclass, field
from typing import Any

from docx import Document
from docx.oxml.ns import qn


@dataclass
class MatchResult:
    """A matched element with its context."""
    element_type: str  # "paragraph", "run", or "block"
    paragraph_index: int
    paragraph: Any  # docx.text.paragraph.Paragraph
    run: Any = None  # docx.text.run.Run (if element_type == "run")
    run_index: int = -1
    text: str = ""

    # Block-level fields
    block_children: list = None  # list of body child elements
    block_start: int = -1  # body child index start
    block_end: int = -1    # body child index end (exclusive)
    block_heading_text: str = ""

    def to_dict(self):
        d = {
            "type": self.element_type,
            "paragraph_index": self.paragraph_index,
            "text": self.text,
        }
        if self.run is not None:
            d["run_index"] = self.run_index
        if self.element_type == "block":
            d["block_start"] = self.block_start
            d["block_end"] = self.block_end
            d["block_heading"] = self.block_heading_text
            d["block_size"] = len(self.block_children) if self.block_children else 0
            # Summarize children types
            if self.block_children:
                tags = []
                for child in self.block_children:
                    tag = child.tag.split("}")[-1]
                    tags.append(tag)
                from collections import Counter
                d["block_children_summary"] = dict(Counter(tags))
        return d


def get_outline_level(para):
    """Get outline level of a paragraph.

    Checks direct paragraph properties first, then falls back to style definition.
    Returns int or None if not set.
    """
    # 1. Direct paragraph property
    ppr = para._element.find(qn("w:pPr"))
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None:
            val = outline.get(qn("w:val"))
            if val is not None:
                return int(val)

    # 2. Style definition (inherited)
    if para.style and para.style.element is not None:
        style_ppr = para.style.element.find(qn("w:pPr"))
        if style_ppr is not None:
            outline = style_ppr.find(qn("w:outlineLvl"))
            if outline is not None:
                val = outline.get(qn("w:val"))
                if val is not None:
                    return int(val)

    return None


def build_heading_map(doc):
    """Build a map: paragraph_index -> nearest preceding heading text.

    Uses style name (Heading X) OR outline level (any level) to detect headings.
    Returns dict of {para_index: heading_text}.
    """
    heading_map = {}
    current_heading = None
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        outline_lvl = get_outline_level(para)

        is_heading = (
            style_name.startswith("Heading")
            or outline_lvl is not None
        )
        if is_heading and para.text.strip():
            current_heading = para.text

        heading_map[i] = current_heading
    return heading_map


def _match_leaf(leaf, para, para_index, heading_map, in_table_para_ids=None):
    """Check if a single paragraph matches a leaf filter.

    Returns list of MatchResult (could be paragraph-level or run-level matches).
    """
    style_name = para.style.name if para.style else ""
    outline_lvl = get_outline_level(para)

    # --- paragraph-level filters ---

    # style
    if "style" in leaf:
        if style_name != leaf["style"]:
            return []

    # outline_level
    if "outline_level" in leaf:
        if outline_lvl != leaf["outline_level"]:
            return []

    # paragraph_range
    if "paragraph_range" in leaf:
        start, end = leaf["paragraph_range"]
        if not (start <= para_index <= end):
            return []

    # in_table
    if "in_table" in leaf:
        is_in_table = para._element.getparent().tag == qn("w:tc") if in_table_para_ids is None else para_index in in_table_para_ids
        if leaf["in_table"] != is_in_table:
            return []

    # under_heading
    if "under_heading" in leaf:
        heading = heading_map.get(para_index)
        if heading is None or leaf["under_heading"] not in heading:
            return []

    # text_contains (paragraph level)
    if "text_contains" in leaf:
        if leaf["text_contains"] not in para.text:
            return []

    # text_regex (paragraph level)
    if "text_regex" in leaf:
        if not re.search(leaf["text_regex"], para.text):
            return []

    # bold (paragraph has any bold run)
    if "bold" in leaf:
        has_bold = any(r.bold for r in para.runs)
        if leaf["bold"] != has_bold:
            return []

    # italic (paragraph has any italic run)
    if "italic" in leaf:
        has_italic = any(r.italic for r in para.runs)
        if leaf["italic"] != has_italic:
            return []

    # font (paragraph has any run with this font)
    if "font" in leaf:
        has_font = any(_run_font(r) == leaf["font"] for r in para.runs)
        if not has_font:
            return []

    # --- run-level filters ---
    run_filters = any(k.startswith("run_") for k in leaf)

    if run_filters:
        # Match individual runs
        results = []
        for ri, run in enumerate(para.runs):
            if _match_run_filters(leaf, run):
                results.append(MatchResult(
                    element_type="run",
                    paragraph_index=para_index,
                    paragraph=para,
                    run=run,
                    run_index=ri,
                    text=run.text,
                ))
        return results

    # No run filters → paragraph-level match
    return [MatchResult(
        element_type="paragraph",
        paragraph_index=para_index,
        paragraph=para,
        text=para.text,
    )]


def _run_font(run):
    """Get effective font name for a run, preferring eastAsia."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            ea = rfonts.get(qn("w:eastAsia"))
            if ea:
                return ea
    return run.font.name or ""


def _match_run_filters(leaf, run):
    """Check if a run matches run_* filters."""
    if "run_text_contains" in leaf:
        if leaf["run_text_contains"] not in run.text:
            return False

    if "run_text_regex" in leaf:
        if not re.search(leaf["run_text_regex"], run.text):
            return False

    if "run_bold" in leaf:
        if leaf["run_bold"] != bool(run.bold):
            return False

    if "run_italic" in leaf:
        if leaf["run_italic"] != bool(run.italic):
            return False

    if "run_font" in leaf:
        if _run_font(run) != leaf["run_font"]:
            return False

    if "run_size" in leaf:
        size = run.font.size
        size_pt = size.pt if size else None
        if size_pt != leaf["run_size"]:
            return False

    return True


def _find_heading_blocks(doc, filter_spec):
    """Find heading blocks in the document body.

    A heading block is all body children from a heading paragraph to the
    next heading at the same or higher outline level / style level.

    Args:
        doc: python-docx Document
        filter_spec: dict with heading_block/heading_block_regex and optional
                     heading_level, heading_outline_level

    Returns:
        List of MatchResult with element_type="block".
    """
    body = doc.element.body
    children = list(body)
    paragraphs = list(doc.paragraphs)

    # Build a map from paragraph._element -> paragraph index
    para_elem_to_idx = {}
    for i, para in enumerate(paragraphs):
        para_elem_to_idx[id(para._element)] = i

    # Determine heading matching criteria
    heading_text = filter_spec.get("heading_block", "")
    heading_regex = filter_spec.get("heading_block_regex", None)
    target_level = filter_spec.get("heading_level")  # style level (1, 2, 3...)
    target_outline = filter_spec.get("heading_outline_level")  # outline level (0, 1, 2...)

    # Build style lookup dict once (avoid O(n*m) per-paragraph lookup)
    style_outline_cache = {}
    for s in doc.styles:
        if s.element is not None:
            s_ppr = s.element.find(qn("w:pPr"))
            if s_ppr is not None:
                s_outline = s_ppr.find(qn("w:outlineLvl"))
                if s_outline is not None:
                    s_val = s_outline.get(qn("w:val"))
                    if s_val is not None:
                        style_outline_cache[s.style_id] = int(s_val)

    def _get_heading_level(child_elem):
        """Get the effective heading level of a body child element.

        Returns (level, text) or None if not a heading.
        Level: lower number = higher heading. Uses outline level if available,
        otherwise style-based heading level.
        """
        tag = child_elem.tag.split("}")[-1]
        if tag != "p":
            return None

        # Get paragraph text
        text = ""
        for t_elem in child_elem.iter(qn("w:t")):
            text += (t_elem.text or "")

        # Check outline level (direct paragraph property)
        ppr = child_elem.find(qn("w:pPr"))
        outline_lvl = None
        if ppr is not None:
            outline = ppr.find(qn("w:outlineLvl"))
            if outline is not None:
                val = outline.get(qn("w:val"))
                if val is not None:
                    outline_lvl = int(val)

        # Check style
        style_lvl = None
        style_id = None
        if ppr is not None:
            pstyle = ppr.find(qn("w:pStyle"))
            if pstyle is not None:
                style_id = pstyle.get(qn("w:val"))
                if style_id and style_id.lower().startswith("heading"):
                    try:
                        style_lvl = int(re.search(r'\d+', style_id).group())
                    except (AttributeError, ValueError):
                        pass

        # If no direct outline level, check style definition (cached)
        if outline_lvl is None and style_id is not None:
            cached = style_outline_cache.get(style_id)
            if cached is not None:
                outline_lvl = cached

        # Use outline level preferentially, fall back to style level
        effective_level = outline_lvl if outline_lvl is not None else style_lvl
        if effective_level is not None:
            return (effective_level, text.strip())

        return None

    def _matches_heading(heading_info):
        """Check if a heading matches the filter criteria."""
        if heading_info is None:
            return False
        level, text = heading_info

        # Check level constraints
        if target_level is not None and level != target_level - 1:
            # target_level is 1-based (1=chapter), outline is 0-based
            # But for heading_level we match style level directly
            # Check if this is a style-based heading with the right level
            return False
        if target_outline is not None and level != target_outline:
            return False

        # Check text match (skip if only outline_level/level specified)
        if heading_regex is not None:
            if not re.search(heading_regex, text):
                return False
        elif heading_text:
            if heading_text not in text:
                return False

        return True

    # Find all headings and their positions in the body children list
    heading_positions = []  # list of (child_index, level, text, para_idx)
    for ci, child in enumerate(children):
        info = _get_heading_level(child)
        if info is not None:
            level, text = info
            para_idx = para_elem_to_idx.get(id(child), -1)
            heading_positions.append((ci, level, text, para_idx))

    # Find matching headings and extract their blocks
    results = []
    for hi, (ci, level, text, para_idx) in enumerate(heading_positions):
        if not _matches_heading((level, text)):
            continue

        # Determine block end: next heading at same or higher level (lower number)
        block_end_ci = len(children)
        for hi2 in range(hi + 1, len(heading_positions)):
            _, next_level, _, _ = heading_positions[hi2]
            if next_level <= level:
                block_end_ci = heading_positions[hi2][0]
                break

        block_children = children[ci:block_end_ci]

        # Find the paragraph text for display
        display_text = text if text else "(empty heading)"

        results.append(MatchResult(
            element_type="block",
            paragraph_index=para_idx,
            paragraph=None,
            text=display_text,
            block_children=block_children,
            block_start=ci,
            block_end=block_end_ci,
            block_heading_text=display_text,
        ))

    return results


def evaluate_filter(filter_spec, paragraphs, heading_map, in_table_para_ids=None, doc=None):
    """Evaluate a filter spec against document paragraphs.

    Returns list of MatchResult.
    """
    # AND: dict
    if isinstance(filter_spec, dict):
        # Handle explicit and/or
        if "and" in filter_spec:
            sub_specs = filter_spec["and"]
            return _intersect(sub_specs, paragraphs, heading_map, in_table_para_ids, doc)
        if "or" in filter_spec:
            sub_specs = filter_spec["or"]
            return _union(sub_specs, paragraphs, heading_map, in_table_para_ids, doc)

        # Check if this is a block-level filter
        is_block = any(k in filter_spec for k in (
            "heading_block", "heading_block_regex", "heading_outline_level", "heading_level"))
        if is_block:
            if doc is None:
                raise ValueError("Block filter requires doc parameter")
            return _find_heading_blocks(doc, filter_spec)

        # Leaf filter: match each paragraph
        results = []
        for i, para in enumerate(paragraphs):
            results.extend(_match_leaf(filter_spec, para, i, heading_map, in_table_para_ids))
        return results

    # OR: list
    if isinstance(filter_spec, list):
        return _union(filter_spec, paragraphs, heading_map, in_table_para_ids, doc)

    raise ValueError(f"Invalid filter spec: {type(filter_spec)}")


def _intersect(sub_specs, paragraphs, heading_map, in_table_para_ids, doc=None):
    """AND: results that appear in ALL sub-specs."""
    if not sub_specs:
        return []

    sets = []
    for spec in sub_specs:
        results = evaluate_filter(spec, paragraphs, heading_map, in_table_para_ids, doc)
        # Use (para_index, run_index, element_type) as key
        keys = {(r.paragraph_index, r.run_index, r.element_type) for r in results}
        sets.append((keys, results))

    # Start with first set's results, keep only those in all sets
    common_keys = sets[0][0]
    for keys, _ in sets[1:]:
        common_keys = common_keys & keys

    # Collect matching results
    all_results = {}
    for keys, results in sets:
        for r in results:
            k = (r.paragraph_index, r.run_index, r.element_type)
            if k in common_keys and k not in all_results:
                all_results[k] = r

    return [all_results[k] for k in sorted(all_results.keys())]


def _union(sub_specs, paragraphs, heading_map, in_table_para_ids, doc=None):
    """OR: results from any sub-spec, deduplicated."""
    seen = set()
    results = []
    for spec in sub_specs:
        for r in evaluate_filter(spec, paragraphs, heading_map, in_table_para_ids, doc):
            key = (r.paragraph_index, r.run_index, r.element_type)
            if key not in seen:
                seen.add(key)
                results.append(r)
    # Sort by position
    results.sort(key=lambda r: (r.paragraph_index, r.run_index))
    return results


def apply_filter(file_path, filter_spec):
    """Apply a filter to a document and return matching results.

    Args:
        file_path: Path to .docx file.
        filter_spec: Filter specification (dict or list).

    Returns:
        List of MatchResult.
    """
    doc = Document(file_path)
    paragraphs = list(doc.paragraphs)
    heading_map = build_heading_map(doc)

    # Build in_table set (O(k) instead of O(n*k))
    para_elem_ids = {id(p._element): i for i, p in enumerate(paragraphs)}
    in_table_para_ids = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pid = para_elem_ids.get(id(para._element))
                    if pid is not None:
                        in_table_para_ids.add(pid)

    return evaluate_filter(filter_spec, paragraphs, heading_map, in_table_para_ids, doc=doc)
