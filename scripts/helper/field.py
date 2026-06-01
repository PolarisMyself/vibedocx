"""Field module - list, update, insert, and refresh document fields."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def enable_auto_update_fields(doc):
    """Enable 'Update Fields on Open' in document settings.

    This tells Word/LibreOffice to automatically recalculate all fields
    when the document is opened. No external dependencies needed.
    """
    settings = doc.settings.element
    # Remove existing updateFields if any
    for uf in settings.findall(qn("w:updateFields")):
        settings.remove(uf)
    # Add new one
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.insert(0, uf)


def _parse_fields(doc):
    """Parse all fields from document body, headers, and footers.

    Returns list of field dicts with: id, type, instruction, location,
    paragraph_index, result_cached, preceding_text, following_text.
    """
    fields = []
    field_id = 0

    # Scan body paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        para_fields = _extract_fields_from_element(
            para._element, location="body", paragraph_index=p_idx
        )
        for f in para_fields:
            f["id"] = field_id
            field_id += 1
            fields.append(f)

    # Scan headers and footers
    for section in doc.sections:
        for hf, loc in [(section.header, "header"), (section.footer, "footer")]:
            if hf is not None:
                for para in hf.paragraphs:
                    para_fields = _extract_fields_from_element(
                        para._element, location=loc, paragraph_index=None
                    )
                    for f in para_fields:
                        f["id"] = field_id
                        field_id += 1
                        fields.append(f)

    return fields


def _extract_fields_from_element(element, location, paragraph_index):
    """Extract fields from a paragraph element by scanning fldChar markers."""
    fields = []
    current_field = None
    instr_parts = []
    result_parts = []
    before_separate = False

    for child in element:
        tag = child.tag

        if tag == qn("w:r"):
            fld_char = child.find(qn("w:fldChar"))
            if fld_char is not None:
                fld_type = fld_char.get(qn("w:fldCharType"))
                if fld_type == "begin":
                    current_field = {
                        "type": "UNKNOWN",
                        "instruction": "",
                        "location": location,
                        "paragraph_index": paragraph_index,
                        "result_cached": "",
                        "preceding_text": "",
                        "following_text": "",
                    }
                    instr_parts = []
                    result_parts = []
                    before_separate = True
                elif fld_type == "separate":
                    if current_field is not None:
                        current_field["instruction"] = " ".join(instr_parts).strip()
                        current_field["type"] = _detect_field_type(current_field["instruction"])
                        before_separate = False
                elif fld_type == "end":
                    if current_field is not None:
                        current_field["result_cached"] = "".join(result_parts)
                        fields.append(current_field)
                        current_field = None
            else:
                # Check for instrText (field instruction)
                for it_elem in child.findall(qn("w:instrText")):
                    text = it_elem.text or ""
                    if current_field is not None and before_separate:
                        instr_parts.append(text)
                # Check for t (field result or other text)
                for t_elem in child.findall(qn("w:t")):
                    text = t_elem.text or ""
                    if current_field is not None and not before_separate:
                        result_parts.append(text)

    return fields


def _detect_field_type(instruction):
    """Detect field type from instruction text."""
    stripped = instruction.strip().upper()
    # Common field types
    known_types = [
        "PAGE", "NUMPAGES", "DATE", "TIME", "TOC", "PAGEREF", "REF",
        "SEQ", "IF", "DOCPROPERTY", "INCLUDEPICTURE", "HYPERLINK",
        "NOTEREF", "FOOTNOTE", "AUTONUM", "STYLEREF",
    ]
    for ft in known_types:
        if stripped.startswith(ft):
            return ft
    return stripped.split()[0] if stripped else "UNKNOWN"


def field_list(file_path, field_type=None):
    """List all fields in the document.

    Args:
        file_path: Path to .docx file.
        field_type: Optional filter (e.g. "PAGE", "DATE", "TOC").

    Returns:
        Dict with "fields" list.
    """
    doc = Document(file_path)
    fields = _parse_fields(doc)

    if field_type:
        fields = [f for f in fields if f["type"] == field_type.upper()]

    return {"fields": fields}


def field_update(file_path, field_id=None, field_type=None, nth=1,
                 instruction=None, output=None):
    """Update a field's instruction.

    Identify field by field_id (exact) or field_type + nth (by type and occurrence).
    """
    doc = Document(file_path)
    fields = _parse_fields(doc)

    # Find target field
    target = None
    if field_id is not None:
        for f in fields:
            if f["id"] == field_id:
                target = f
                break
        if target is None:
            raise ValueError(f"Field with id {field_id} not found")
    elif field_type is not None:
        matches = [f for f in fields if f["type"] == field_type.upper()]
        if nth <= len(matches):
            target = matches[nth - 1]
        else:
            raise ValueError(f"Field {field_type} #{nth} not found (only {len(matches)} exist)")
    else:
        raise ValueError("Specify field_id or field_type")

    # Update the instruction in the XML
    _update_field_instruction(doc, target, instruction)
    enable_auto_update_fields(doc)
    doc.save(output)


def _update_field_instruction(doc, field_info, new_instruction):
    """Update a specific field's instruction text in the XML."""
    para = _find_paragraph(doc, field_info)
    if para is None:
        return

    # Find the instrText in the paragraph
    in_field = False
    found_separate = False

    for child in para:
        if child.tag != qn("w:r"):
            continue

        fld_char = child.find(qn("w:fldChar"))
        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_field = True
            elif fld_type == "separate" and in_field:
                found_separate = True
            elif fld_type == "end" and in_field:
                break
        elif in_field and not found_separate:
            instr = child.find(qn("w:instrText"))
            if instr is not None:
                instr.text = new_instruction


def _find_paragraph(doc, field_info):
    """Find the paragraph element containing a field.

    Returns the lxml element (CT_P), not a Paragraph object.
    """
    location = field_info["location"]
    p_idx = field_info["paragraph_index"]

    if location == "body":
        if p_idx is not None and p_idx < len(doc.paragraphs):
            return doc.paragraphs[p_idx]._element
    elif location in ("header", "footer"):
        for section in doc.sections:
            hf = section.header if location == "header" else section.footer
            if hf is not None:
                for i, para in enumerate(hf.paragraphs):
                    if i == p_idx:
                        return para._element
    return None


def field_insert(file_path, after_paragraph=None, field_type=None, fmt=None,
                 location='body', output=None):
    """Insert a new field into the document.

    Args:
        file_path: Input .docx path.
        after_paragraph: Insert after this paragraph index (for location='body').
        field_type: Field type (PAGE, NUMPAGES, DATE, TOC, etc.).
        format: Optional field format/switch.
        location: 'body', 'header', or 'footer'.
        output: Output file path.
    """
    doc = Document(file_path)

    # Build field instruction
    ft = field_type.upper()
    if ft == "PAGE":
        instr = " PAGE "
        if fmt:
            instr += f"\\* {fmt.upper()} "
        result = ""
    elif ft == "NUMPAGES":
        instr = " NUMPAGES "
        result = ""
    elif ft == "DATE":
        instr = " DATE "
        if fmt:
            instr += f'\\@ "{fmt}" '
        from datetime import datetime
        result = datetime.now().strftime(fmt if fmt else "%Y-%m-%d")
    elif ft == "TOC":
        levels = fmt or "1-3"
        instr = f' TOC \\o "{levels}" \\h \\z \\u '
        result = "[目录]"
    else:
        instr = f" {ft} "
        if fmt:
            instr += fmt + " "
        result = ""

    # Build field paragraph
    new_para = OxmlElement("w:p")

    # begin
    r1 = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)
    new_para.append(r1)

    # instruction
    r2 = OxmlElement("w:r")
    instr_elem = OxmlElement("w:instrText")
    instr_elem.set(qn("xml:space"), "preserve")
    instr_elem.text = instr
    r2.append(instr_elem)
    new_para.append(r2)

    # separate
    r3 = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fld_sep)
    new_para.append(r3)

    # result
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = result
    r4.append(t)
    new_para.append(r4)

    # end
    r5 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5.append(fld_end)
    new_para.append(r5)

    if location in ('header', 'footer'):
        section = doc.sections[0]
        hf = section.header if location == 'header' else section.footer
        hf.is_linked_to_previous = False
        hf._element.append(new_para)
        if after_paragraph is not None:
            hf_paras = list(hf.paragraphs)
            if 0 <= after_paragraph < len(hf_paras):
                hf._element.remove(new_para)
                hf_paras[after_paragraph]._element.addnext(new_para)
    else:
        paras = list(doc.paragraphs)
        if after_paragraph is None or after_paragraph < 0 or after_paragraph >= len(paras):
            raise ValueError(f"Paragraph index {after_paragraph} out of range (0-{len(paras)-1})")
        target_para = paras[after_paragraph]._element
        parent = target_para.getparent()
        target_idx = list(parent).index(target_para)
        parent.insert(target_idx + 1, new_para)

    enable_auto_update_fields(doc)
    doc.save(output)


def field_refresh(file_path, field_type=None, output=None):
    """Refresh cached results of fields.

    For computable fields (PAGE, NUMPAGES, DATE), updates the cached result text.
    TOC and other complex fields can only have their instruction updated;
    full refresh requires opening in Word.

    Returns count of fields refreshed.
    """
    doc = Document(file_path)
    fields = _parse_fields(doc)

    if field_type:
        fields = [f for f in fields if f["type"] == field_type.upper()]

    count = 0
    for f in fields:
        ft = f["type"]
        new_result = None

        if ft == "PAGE":
            new_result = ""  # Cannot determine actual page number; clear cache
        elif ft == "NUMPAGES":
            new_result = ""  # Same limitation
        elif ft == "DATE":
            from datetime import datetime
            # Try to extract format from instruction
            instr = f["instruction"]
            if '\\@' in instr:
                # Keep the instruction as-is, update result to current date
                new_result = datetime.now().strftime("%Y-%m-%d")
            else:
                new_result = datetime.now().strftime("%Y-%m-%d")

        if new_result is not None:
            _update_field_result(doc, f, new_result)
            count += 1

    enable_auto_update_fields(doc)
    doc.save(output)
    return count


def _update_field_result(doc, field_info, new_result):
    """Update a field's cached result text."""
    para = _find_paragraph(doc, field_info)
    if para is None:
        return

    in_field = False
    found_separate = False
    result_updated = False

    for child in para:
        if child.tag != qn("w:r"):
            continue

        fld_char = child.find(qn("w:fldChar"))
        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_field = True
            elif fld_type == "separate" and in_field:
                found_separate = True
            elif fld_type == "end" and in_field:
                break
        elif in_field and found_separate and not result_updated:
            for t_elem in child.findall(qn("w:t")):
                t_elem.text = new_result
                result_updated = True
                break
