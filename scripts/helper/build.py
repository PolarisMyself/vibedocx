"""Create module - generate new DOCX documents from scratch."""

import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from helper.config import PAPER_SIZES, ALIGN_MAP
from helper.units import parse_size


def _set_run_font(run, font_cn=None, font_en=None, size=None, bold=None, italic=None):
    """Set run font properties."""
    from docx.oxml import OxmlElement
    if size is not None:
        run.font.size = parse_size(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if font_cn or font_en:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if font_cn:
            rFonts.set(qn('w:eastAsia'), font_cn)
        if font_en:
            rFonts.set(qn('w:ascii'), font_en)
            rFonts.set(qn('w:hAnsi'), font_en)
        if font_cn and not font_en:
            rFonts.set(qn('w:ascii'), font_cn)
            rFonts.set(qn('w:hAnsi'), font_cn)


def _apply_styles(doc, styles_def):
    """Apply style definitions to the document.

    styles_def: dict of style_name -> properties dict.
    Properties: font, font_en, size, bold, italic, align,
                line_spacing, space_before, space_after, indent_first
    """
    for style_name, props in styles_def.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue  # skip unknown styles

        # Font
        if 'font' in props or 'font_en' in props:
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rfonts)
            if 'font' in props:
                rfonts.set(qn("w:eastAsia"), props['font'])
                style.font.name = props['font']
            if 'font_en' in props:
                rfonts.set(qn("w:ascii"), props['font_en'])
                rfonts.set(qn("w:hAnsi"), props['font_en'])
                style.font.name = props['font_en']

        if 'size' in props:
            style.font.size = parse_size(props['size'])
        if 'bold' in props:
            style.font.bold = props['bold']
        if 'italic' in props:
            style.font.italic = props['italic']

        # Paragraph formatting (only for paragraph styles)
        if style.type == 1:
            pf = style.paragraph_format
            if 'align' in props:
                pf.alignment = ALIGN_MAP.get(props['align'])
            if 'line_spacing' in props:
                pf.line_spacing = props['line_spacing']
            if 'space_before' in props:
                pf.space_before = Pt(props['space_before'])
            if 'space_after' in props:
                pf.space_after = Pt(props['space_after'])
            if 'indent_first' in props:
                # Convert char count to points (approximate)
                fs = props.get('size', 12)
                pf.first_line_indent = Pt(props['indent_first'] * fs * 0.35)


def _add_page_break(doc):
    """Insert a page break."""
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def _add_section_break(doc, break_type='nextPage'):
    """Insert a section break.

    break_type: 'nextPage', 'continuous', 'evenPage', 'oddPage'
    """
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    doc_sect = doc.sections[-1]
    new_sectPr = OxmlElement('w:sectPr')
    type_elem = OxmlElement('w:type')
    type_elem.set(qn('w:val'), break_type)
    new_sectPr.append(type_elem)
    # Copy page setup from the current section
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(int(doc_sect.page_width.emu)))
    pgSz.set(qn('w:h'), str(int(doc_sect.page_height.emu)))
    new_sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), str(int(doc_sect.top_margin.emu)))
    pgMar.set(qn('w:bottom'), str(int(doc_sect.bottom_margin.emu)))
    pgMar.set(qn('w:left'), str(int(doc_sect.left_margin.emu)))
    pgMar.set(qn('w:right'), str(int(doc_sect.right_margin.emu)))
    new_sectPr.append(pgMar)
    pPr.append(new_sectPr)


def _add_table(doc, table_def):
    """Add a table from a definition dict.

    table_def keys: caption, headers, rows, width_pct
    """
    caption = table_def.get('caption')
    headers = table_def.get('headers', [])
    rows = table_def.get('rows', [])

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(caption)
        _set_run_font(run, font_cn='黑体', size=10.5, bold=True)

    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    num_rows = (1 if headers else 0) + len(rows)
    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    row_idx = 0
    # Header row
    if headers:
        for j, text in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ''
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(str(text))
            _set_run_font(run, font_cn='黑体', size=10.5, bold=True)
        # Shade header
        for cell in table.rows[0].cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)
        row_idx = 1

    # Data rows
    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            cell = table.rows[row_idx + i].cells[j]
            cell.text = ''
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(str(text))
            _set_run_font(run, font_cn='宋体', size=10.5)


def create_blank(output, paper='A4', margin_top=2.54, margin_bottom=2.54,
                 margin_left=3.17, margin_right=3.17, orientation='portrait'):
    """Create a blank DOCX with page setup.

    Returns the output path.
    """
    doc = Document()
    section = doc.sections[0]

    # Paper
    w, h = PAPER_SIZES.get(paper.upper(), (21.0, 29.7))
    section.page_width = Cm(w)
    section.page_height = Cm(h)

    # Margins
    section.top_margin = Cm(margin_top)
    section.bottom_margin = Cm(margin_bottom)
    section.left_margin = Cm(margin_left)
    section.right_margin = Cm(margin_right)

    # Orientation
    if orientation == 'landscape':
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

    doc.save(output)
    return output


def create_structure(output, structure_json):
    """Create a document from a JSON structure definition.

    structure_json can be a dict or a JSON string.

    Supported keys:
      - page: {paper, margin_top, margin_bottom, margin_left, margin_right, orientation}
      - styles: {style_name: {font, font_en, size, bold, italic, align, line_spacing, ...}}
      - header_text / footer_text: str
      - content: list of content items, each is a dict with one of:
          {heading: level, text: "..."}
          {body: "...", indent: bool, align: "..."}
          {page_break: true}
          {section_break: "nextPage|continuous"}
          {table: {caption, headers, rows}}
          {blank_lines: n}
    """
    if isinstance(structure_json, str):
        structure_json = json.loads(structure_json)

    doc = Document()

    # ── Page setup ──
    page = structure_json.get('page', {})
    section = doc.sections[0]
    paper = page.get('paper', 'A4')
    w, h = PAPER_SIZES.get(paper.upper(), (21.0, 29.7))
    section.page_width = Cm(w)
    section.page_height = Cm(h)
    section.top_margin = Cm(page.get('margin_top', 2.54))
    section.bottom_margin = Cm(page.get('margin_bottom', 2.54))
    section.left_margin = Cm(page.get('margin_left', 3.17))
    section.right_margin = Cm(page.get('margin_right', 3.17))
    if page.get('orientation') == 'landscape':
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

    # ── Styles ──
    styles_def = structure_json.get('styles', {})
    if styles_def:
        _apply_styles(doc, styles_def)

    # ── Header / Footer ──
    if 'header_text' in structure_json:
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            header.paragraphs[0].text = structure_json['header_text']
        else:
            header.add_paragraph(structure_json['header_text'])

    if 'footer_text' in structure_json:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            footer.paragraphs[0].text = structure_json['footer_text']
        else:
            footer.add_paragraph(structure_json['footer_text'])

    # ── Content ──
    for item in structure_json.get('content', []):
        if 'heading' in item:
            level = item['heading']
            text = item.get('text', '')
            p = doc.add_heading(text, level=level)

        elif 'body' in item:
            text = item['body']
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = item.get('line_spacing', 1.5)
            align = item.get('align')
            if align:
                p.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if item.get('indent', True):
                p.paragraph_format.first_line_indent = Pt(24)
            run = p.add_run(text)
            font_cn = item.get('font_cn', '宋体')
            font_en = item.get('font_en', 'Times New Roman')
            size = item.get('size', 12)
            bold = item.get('bold')
            italic = item.get('italic')
            _set_run_font(run, font_cn=font_cn, font_en=font_en, size=size,
                          bold=bold, italic=italic)

        elif 'page_break' in item and item['page_break']:
            _add_page_break(doc)

        elif 'section_break' in item:
            _add_section_break(doc, item.get('section_break', 'nextPage'))

        elif 'table' in item:
            _add_table(doc, item['table'])

        elif 'blank_lines' in item:
            for _ in range(item['blank_lines']):
                doc.add_paragraph()

    doc.save(output)
    return output
