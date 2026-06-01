"""Numbering module - chapter number updating after block operations.

Handles:
- Section numbers: 3.1, 3.2.1, etc.
- Figure/table references: 图3-1, 表3-2, Figure 3-1, Table 3-1, etc.
- Chapter names in overview sections
- Cross-run text (patterns spanning multiple runs)
"""

import re
from docx import Document
from docx.oxml.ns import qn


# ──────────────────────────────────────────────
# Default pattern sets
# ──────────────────────────────────────────────

# Chinese academic document patterns
PATTERNS_CN = {
    "name": "cn",
    "description": "中文学术文档",
    "section": r"\b{ch}\.\d+(?:\.\d+)*\b",
    "refs": [
        # 图/表 + 章节号-序号 (various separators)
        {"prefix": "图", "sep": ["-", "－", "‐", "–"], "pattern": r"图\s*{ch}\s*[{sep}]\s*\d+"},
        {"prefix": "表", "sep": ["-", "－", "‐", "–"], "pattern": r"表\s*{ch}\s*[{sep}]\s*\d+"},
        # 带空格的变体
        {"prefix": "图", "sep": [" "], "pattern": r"图\s+{ch}\s+{sep}\s*\d+"},
        {"prefix": "表", "sep": [" "], "pattern": r"表\s+{ch}\s+{sep}\s*\d+"},
    ],
    "chapter_name": r"第[三四五六七八九十]+章",
}

# English academic document patterns
PATTERNS_EN = {
    "name": "en",
    "description": "English academic documents",
    "section": r"\b{ch}\.\d+(?:\.\d+)*\b",
    "refs": [
        {"prefix": "Figure", "sep": ["-", " "], "pattern": r"Figure\s*{ch}\s*[{sep}]\s*\d+"},
        {"prefix": "Fig.", "sep": ["-", " "], "pattern": r"Fig\.?\s*{ch}\s*[{sep}]\s*\d+"},
        {"prefix": "Table", "sep": ["-", " "], "pattern": r"Table\s*{ch}\s*[{sep}]\s*\d+"},
        {"prefix": "Chapter", "sep": ["-", " "], "pattern": r"Chapter\s*{ch}\s*[{sep}]\s*\d+"},
    ],
    "chapter_name": r"Chapter\s+\d+",
}

# Combined defaults
DEFAULT_PATTERNS = [PATTERNS_CN, PATTERNS_EN]


def _build_swap_regexes(pattern_def, ch_a, ch_b):
    """Build regex list and replacement functions for a pattern set.

    Returns list of (compiled_regex, replace_func) tuples.
    """
    results = []
    a, b = str(ch_a), str(ch_b)

    # Section numbers: {ch}.1, {ch}.2.1, etc.
    sec_pat = pattern_def["section"].format(ch=f"[{a}{b}]")
    sec_re = re.compile(sec_pat)

    def sec_replace(m):
        num = m.group(0)
        parts = num.split(".")
        if parts[0] == a:
            parts[0] = b
        elif parts[0] == b:
            parts[0] = a
        return ".".join(parts)

    results.append((sec_re, sec_replace))

    # Figure/table references
    for ref in pattern_def.get("refs", []):
        seps = ref.get("sep", ["-"])
        sep_char_class = "".join(re.escape(s) for s in seps)
        pat_str = ref["pattern"].format(ch=f"[{a}{b}]", sep=sep_char_class)
        try:
            ref_re = re.compile(pat_str)
        except re.error:
            continue

        def make_ref_replace(seps_list):
            def ref_replace(m):
                text = m.group(0)
                for sep in seps_list:
                    old_a = f"{a}{sep}"
                    old_b = f"{b}{sep}"
                    if old_a in text:
                        return text.replace(old_a, f"{b}{sep}")
                    if old_b in text:
                        return text.replace(old_b, f"{a}{sep}")
                return text
            return ref_replace

        results.append((ref_re, make_ref_replace(seps)))

    return results


def _swap_in_text(text, regexes):
    """Apply all regex swaps to text. Returns (new_text, changed)."""
    original = text
    for regex, replace_fn in regexes:
        text = regex.sub(replace_fn, text)
    return text, text != original


def _swap_in_paragraph(para, regexes):
    """Swap chapter numbers across runs within a paragraph.

    Handles cross-run patterns by reconstructing full text.
    Handles length changes (e.g. 9→10) by merging overflow into last run.
    Returns count of modified runs.
    """
    if not para.runs:
        return 0

    full_text = "".join(r.text for r in para.runs)
    new_text, changed = _swap_in_text(full_text, regexes)
    if not changed:
        return 0

    # Redistribute new_text back to runs.
    # If length changed, distribute proportionally — last run takes remainder.
    count = 0
    runs = para.runs
    run_texts = [r.text for r in runs]
    total_original_len = sum(len(t) for t in run_texts)
    total_new_len = len(new_text)

    pos = 0
    for ri, run in enumerate(runs):
        if ri < len(runs) - 1:
            ratio = len(run_texts[ri]) / total_original_len if total_original_len else 0
            run_len = round(total_new_len * ratio)
            new_run_text = new_text[pos:pos + run_len]
        else:
            new_run_text = new_text[pos:]
        if new_run_text != run.text:
            run.text = new_run_text
            count += 1
        pos += len(new_run_text)

    return count


def _swap_chapter_names(text, ch_a, ch_b, chapter_names):
    """Swap chapter name references (e.g., '第三章 标题' <-> '第四章 标题').

    Args:
        text: Input text.
        ch_a, ch_b: Chapter numbers to swap.
        chapter_names: Dict of {chapter_number: chapter_title} or None.

    Returns:
        (new_text, changed)
    """
    if not chapter_names:
        return text, False

    a, b = str(ch_a), str(ch_b)
    name_a = chapter_names.get(a)
    name_b = chapter_names.get(b)
    if not name_a or not name_b:
        return text, False

    # Build patterns: 第三章 标题
    _NUM_TO_CN = {
        "1": "一", "2": "二", "3": "三", "4": "四",
        "5": "五", "6": "六", "7": "七", "8": "八",
        "9": "九", "10": "十", "11": "十一", "12": "十二",
        "13": "十三", "14": "十四", "15": "十五", "16": "十六",
        "17": "十七", "18": "十八", "19": "十九", "20": "二十",
        "21": "二十一", "22": "二十二", "23": "二十三", "24": "二十四",
        "25": "二十五", "26": "二十六", "27": "二十七", "28": "二十八",
        "29": "二十九", "30": "三十", "31": "三十一", "32": "三十二",
        "33": "三十三", "34": "三十四", "35": "三十五", "36": "三十六",
    }
    cn_a = _NUM_TO_CN.get(a, a)
    cn_b = _NUM_TO_CN.get(b, b)

    pat_a = f"第{cn_a}章\\s*{re.escape(name_a)}"
    pat_b = f"第{cn_b}章\\s*{re.escape(name_b)}"

    new_text = text
    new_text = re.sub(pat_a, f"___CH_NAME_TMP_B___", new_text)
    new_text = re.sub(pat_b, f"第{cn_a}章 {name_b}", new_text)
    new_text = new_text.replace("___CH_NAME_TMP_B___", f"第{cn_b}章 {name_a}")

    return new_text, new_text != text


# ──────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────


def update_chapter_numbers(
    file_path,
    ch_a,
    ch_b,
    chapter_names=None,
    patterns=None,
    output=None,
):
    """Update chapter numbers after a chapter swap.

    Args:
        file_path: Path to .docx file.
        ch_a: First chapter number (int).
        ch_b: Second chapter number (int).
        chapter_names: Optional dict of {str(ch_num): chapter_title} for name swaps.
            e.g. {"3": "可视化平台需求分析", "4": "可视化平台总体设计"}
        patterns: List of pattern dicts to use. Defaults to DEFAULT_PATTERNS.
            Each dict must have "section" and "refs" keys (see PATTERNS_CN).
        output: Output file path.

    Returns:
        Total count of modified runs.
    """
    if patterns is None:
        patterns = DEFAULT_PATTERNS

    doc = Document(file_path)

    # Build all regexes from all pattern sets, deduplicating
    seen_patterns = set()
    all_regexes = []
    for pat_def in patterns:
        for regex, replace_fn in _build_swap_regexes(pat_def, ch_a, ch_b):
            if regex.pattern not in seen_patterns:
                seen_patterns.add(regex.pattern)
                all_regexes.append((regex, replace_fn))

    total = 0

    # Process paragraphs
    for para in doc.paragraphs:
        # Swap section numbers and refs
        total += _swap_in_paragraph(para, all_regexes)
        # Swap chapter names if applicable
        if chapter_names:
            for run in para.runs:
                new_text, changed = _swap_chapter_names(
                    run.text, ch_a, ch_b, chapter_names
                )
                if changed:
                    run.text = new_text
                    total += 1

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += _swap_in_paragraph(para, all_regexes)
                    if chapter_names:
                        for run in para.runs:
                            new_text, changed = _swap_chapter_names(
                                run.text, ch_a, ch_b, chapter_names
                            )
                            if changed:
                                run.text = new_text
                                total += 1

    # Process headers/footers
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    total += _swap_in_paragraph(para, all_regexes)

    doc.save(output)
    return total


# ──────────────────────────────────────────────
# Figure / Table auto-numbering
# ──────────────────────────────────────────────

FIGTAB_NS = 'urn:opendocx:refs'


def _load_figtab(doc):
    """Load figure/table metadata from document settings."""
    settings = doc.settings.element
    elem = settings.find(f'{{{FIGTAB_NS}}}figtab')
    if elem is not None and elem.text:
        import json
        try:
            return json.loads(elem.text)
        except json.JSONDecodeError:
            pass
    return {}


def _save_figtab(doc, data):
    """Save figure/table metadata to document settings."""
    from lxml import etree
    import json
    settings = doc.settings.element
    for old in settings.findall(f'{{{FIGTAB_NS}}}figtab'):
        settings.remove(old)
    elem = etree.SubElement(settings, f'{{{FIGTAB_NS}}}figtab')
    elem.text = json.dumps(data, ensure_ascii=False)


def mark_caption_with_ref(run_elem, ref_id, fig_type):
    """Add a figtab marker to a caption run's rPr.

    Args:
        run_elem: The w:r lxml element.
        ref_id: e.g. 'fig:arch' or 'tab:compare'.
        fig_type: 'figure' or 'table'.
    """
    from lxml import etree
    rPr = run_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(run_elem, qn('w:rPr'))
        run_elem.insert(0, rPr)
    marker = etree.SubElement(rPr, f'{{{FIGTAB_NS}}}figtab')
    marker.set('id', ref_id)
    marker.set('type', fig_type)


def update_figure_numbers(file_path, output=None):
    """Scan document for figure/table markers, reassign chapter-based numbering.

    Figures: 图{chapter}-{seq}, Tables: 表{chapter}-{seq}
    Within each chapter, figures and tables are numbered independently.
    Captions are updated in-place.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    figtab = _load_figtab(doc)
    if not figtab:
        doc.save(output)
        return 0

    # Build chapter map: paragraph_index -> (chapter_number, heading_level)
    chapter_map = {}
    current_ch = 1
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        ppr = para._element.find(qn('w:pPr'))
        outline_lvl = None
        if ppr is not None:
            ol = ppr.find(qn('w:outlineLvl'))
            if ol is not None and ol.get(qn('w:val')) is not None:
                outline_lvl = int(ol.get(qn('w:val')))

        is_chapter = False
        if outline_lvl is not None and outline_lvl <= 1:
            is_chapter = True
        elif style_name.startswith('Heading') and '1' in style_name:
            is_chapter = True

        if is_chapter and para.text.strip():
            # Try to extract chapter number
            import re
            m = re.search(r'第([一二三四五六七八九十\d]+)章', para.text)
            if m:
                cn = m.group(1)
                current_ch = _parse_chinese_number(cn) if cn else current_ch
            elif outline_lvl == 0:
                current_ch += 1
        chapter_map[i] = current_ch

    # Assign chapter to each paragraph
    def _get_chapter(para_index):
        if para_index in chapter_map:
            return chapter_map[para_index]
        prev = 1
        for pi in sorted(chapter_map.keys()):
            if pi > para_index:
                break
            prev = chapter_map[pi]
        return prev

    # Scan for figtab markers and group by chapter
    from collections import defaultdict
    ch_figures = defaultdict(list)  # ch -> [(para_index, run, ref_id, old_text)]
    ch_tables = defaultdict(list)

    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                continue
            marker = rPr.find(f'{{{FIGTAB_NS}}}figtab')
            if marker is None:
                continue
            ref_id = marker.get('id')
            ft_type = marker.get('type')
            if ref_id not in figtab:
                continue
            ch = _get_chapter(i)
            if ft_type == 'figure':
                ch_figures[ch].append((i, run, ref_id))
            else:
                ch_tables[ch].append((i, run, ref_id))

    # Generate new numbers and update captions
    count = 0
    # Sort by position within each chapter, assign sequential numbers
    for ch in sorted(set(list(ch_figures.keys()) + list(ch_tables.keys()))):
        figs = sorted(ch_figures.get(ch, []), key=lambda x: x[0])
        for seq, (pi, run, ref_id) in enumerate(figs, 1):
            caption_base = figtab[ref_id].get('caption', '')
            new_text = f'图{ch}-{seq} {caption_base}'
            if run.text != new_text:
                run.text = new_text
                count += 1

        tabs = sorted(ch_tables.get(ch, []), key=lambda x: x[0])
        for seq, (pi, run, ref_id) in enumerate(tabs, 1):
            caption_base = figtab[ref_id].get('caption', '')
            new_text = f'表{ch}-{seq} {caption_base}'
            if run.text != new_text:
                run.text = new_text
                count += 1

    doc.save(output)
    return count


def _parse_chinese_number(s):
    """Parse a Chinese numeral string (e.g. 十一, 二十五) to int.

    Handles 一 through 九十九. Returns int, or None on failure.
    """
    if not s:
        return None
    if s.isdigit():
        return int(s)
    digit_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                 '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
    s = s.strip()
    if s in digit_map:
        return digit_map[s]
    if s == '十':
        return 10
    if s.startswith('十'):
        return 10 + digit_map.get(s[1], 0)
    if s.endswith('十'):
        return digit_map.get(s[0], 0) * 10
    if '十' in s:
        parts = s.split('十', 1)
        tens = digit_map.get(parts[0], 0)
        ones = digit_map.get(parts[1], 0) if parts[1] else 0
        return tens * 10 + ones
    return None
